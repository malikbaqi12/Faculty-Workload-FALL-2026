"""
app.py — Faculty Workload Allocation and Course Recommendation System

Single-file Streamlit application. All backend logic (data loading,
Computing Core Course detection, faculty teaching history, content-based
suitability scoring, constraint-based optimization, and PDF report
generation) lives in this file, organized into clearly labeled sections
below. The companion notebook (faculty_workload_backend.ipynb) contains
the same pipeline for offline experimentation / Google Colab use.

Run with:
    streamlit run app.py
"""

from __future__ import annotations

import io
import re
import traceback
from dataclasses import dataclass, field
from typing import Optional

import numpy as np
import pandas as pd
import pulp
import streamlit as st
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak,
)

# =============================================================================
# SECTION 1 — DATA PREPROCESSING
# (workload / scheme-of-studies loaders, Computing Core Course detection)
# =============================================================================

COMPUTING_CORE_COURSES = {
    "Programming Fundamentals": ["programming fundamental", "programming fundamentals", "intro to computer", "introduction to computing"],
    "Object-Oriented Programming": ["object oriented programming", "oop"],
    "Data Structures": ["data structure", "data structures", "ds"],
    "Computer Networks": ["computer network", "computer networks"],
    "Database Management Systems": ["database system", "database systems", "database management", "dbms"],
    "Analysis of Algorithms": ["analysis of algorithm", "analysis of algorithms", "aoa"],
    "Software Engineering": ["software engineering", "se"],
}
CORE_ALIAS_LOOKUP = {
    alias: canonical
    for canonical, aliases in COMPUTING_CORE_COURSES.items()
    for alias in aliases
}

MAX_CREDIT_HOURS = 9
MAX_THEORY_COURSES = 3
MAX_LAB_COURSES = 1
MAX_PREVIOUSLY_TAUGHT = 1
MAX_FACULTY_PER_COURSE = 2  # duplicate/comparative recommendation cap


def _norm(s: object) -> str:
    if s is None or (isinstance(s, float) and pd.isna(s)):
        return ""
    s = str(s).lower().strip()
    s = re.sub(r"[^a-z0-9 ]", " ", s)
    s = re.sub(r"\s+", " ", s).strip()
    return s


def is_core_course(course_name: str) -> Optional[str]:
    """Return the canonical Computing Core Course name if this course is one, else None."""
    n = _norm(course_name)
    if not n:
        return None
    # Longest aliases first, matched on word boundaries, so short aliases
    # (e.g. "se", "ds") can't false-match inside unrelated words like "security".
    for alias in sorted(CORE_ALIAS_LOOKUP, key=len, reverse=True):
        if re.search(rf"\b{re.escape(alias)}\b", n):
            return CORE_ALIAS_LOOKUP[alias]
    return None


def is_lab(theory_lab_value: str) -> bool:
    return "lab" in _norm(theory_lab_value)


def _clean_text(s: object) -> str:
    """
    Collapse embedded newlines/extra whitespace from wrapped Excel cells
    (e.g. "Ms. Samra\\nJamil" -> "Ms. Samra Jamil") so the same faculty
    member or course isn't accidentally treated as two different ones just
    because of how a cell happened to wrap.
    """
    if s is None or (isinstance(s, float) and pd.isna(s)):
        return ""
    return re.sub(r"\s+", " ", str(s).replace("\n", " ").replace("\r", " ")).strip()


WORKLOAD_COLUMN_ALIASES = {
    "faculty_name": ["faculty name", "faculty", "instructor", "teacher"],
    "course_code": ["course code", "code"],
    "course_name": ["course name", "course", "subject"],
    "theory_lab": ["theory lab", "theory/lab", "type", "th lab"],
    "semester": ["semester", "sem"],
    "credit_hours": ["cr hr", "crhr", "credit hours", "credit hr", "cr h", "credit"],
    "contact_hours": ["cn hr", "cnhr", "contact hours", "contact hr", "con hr", "con h"],
    "program": ["program", "programme", "degree"],
}

SCHEME_COLUMN_ALIASES = {
    "course_code": ["course code", "code"],
    "course_name": ["course name", "course", "subject", "course title"],
    "credit_hours": ["credit hours", "cr hr", "crhr", "credit hr", "credit"],
    "theory_lab": ["theory lab", "theory/lab", "type"],
    "semester": ["semester", "sem"],
    "program": ["program", "programme", "degree"],
}


def _match_columns(columns, alias_map: dict) -> dict:
    norm_cols = {_norm(c): c for c in columns}
    result = {}
    for canonical, aliases in alias_map.items():
        for alias in aliases:
            if alias in norm_cols:
                result[canonical] = norm_cols[alias]
                break
        else:
            for norm_c, orig_c in norm_cols.items():
                if any(a in norm_c for a in aliases):
                    result[canonical] = orig_c
                    break
    return result


def _find_header_row(raw_no_header: pd.DataFrame, alias_map: dict, required_keys: list, max_scan: int = 10) -> int:
    """
    Real-world exports (like a Word table pasted into Excel) often have one or
    more title rows above the actual column headers, e.g.:
        Row 0: "AI Programme Workload Spring 2026"
        Row 1: "Faculty Add"
        Row 2: "Faculty Name | Course Code | Course Name | ..."   <- real header
    Reading with header=0 in that case silently produces a DataFrame with the
    real headers as data and no recognizable columns, which downstream ends up
    as an empty, hard-to-diagnose result. This scans the first few rows and
    picks the one that best matches the expected column aliases.
    """
    best_row, best_score = 0, -1
    scan_limit = min(max_scan, len(raw_no_header))
    for i in range(scan_limit):
        norm_vals = {_norm(v) for v in raw_no_header.iloc[i].tolist()}
        score = 0
        for key in required_keys:
            if any(any(alias == nv or alias in nv for nv in norm_vals) for alias in alias_map[key]):
                score += 1
        if score > best_score:
            best_score = score
            best_row = i
    return best_row


def _read_excel_with_smart_header(file_path_or_buffer, alias_map: dict, required_keys: list) -> pd.DataFrame:
    """Read an Excel file, auto-detecting which row is the real header row."""
    if hasattr(file_path_or_buffer, "seek"):
        file_path_or_buffer.seek(0)
    raw_no_header = pd.read_excel(file_path_or_buffer, header=None)

    header_row = _find_header_row(raw_no_header, alias_map, required_keys)

    if hasattr(file_path_or_buffer, "seek"):
        file_path_or_buffer.seek(0)
    return pd.read_excel(file_path_or_buffer, header=header_row)


def _read_docx_table(file_path_or_buffer, alias_map: dict, required_keys: list) -> pd.DataFrame:
    """
    Extract a table from a Word (.docx) document, such as the department's
    workload table pasted straight into Word. Scans every table in the
    document and every row near the top of each table to find the row that
    best matches the expected column headers (handles a bold title
    paragraph above the table, or a table with its own header a row or two
    down), then builds a DataFrame from everything below that header row.
    """
    import docx

    if hasattr(file_path_or_buffer, "seek"):
        file_path_or_buffer.seek(0)
    document = docx.Document(file_path_or_buffer)

    if not document.tables:
        raise ValueError("No tables were found in the Word document.")

    best_table, best_header_row, best_score = None, 0, -1
    for table in document.tables:
        scan_limit = min(5, len(table.rows))
        for i in range(scan_limit):
            norm_vals = {_norm(cell.text) for cell in table.rows[i].cells}
            score = 0
            for key in required_keys:
                if any(any(alias == nv or alias in nv for nv in norm_vals) for alias in alias_map[key]):
                    score += 1
            if score > best_score:
                best_score = score
                best_table = table
                best_header_row = i

    if best_table is None or best_score <= 0:
        raise ValueError(
            "Could not find a table with recognizable column headers "
            "(e.g. Faculty Name / Course Name) anywhere in the Word document."
        )

    header_cells = [cell.text.strip() for cell in best_table.rows[best_header_row].cells]
    data_rows = []
    for row in best_table.rows[best_header_row + 1:]:
        cells = [cell.text for cell in row.cells]
        # Section-divider rows (e.g. "CS Merge Labs", "Combine") are stored by
        # python-docx as a horizontally-merged cell repeated across every
        # column — i.e. every non-empty cell in the row has the same text.
        # Those are sub-headings in the source table, not real faculty/course
        # data, so drop them here rather than let them masquerade as a course.
        non_empty = {c.strip() for c in cells if c.strip()}
        if len(non_empty) <= 1:
            continue
        data_rows.append(cells)

    return pd.DataFrame(data_rows, columns=header_cells)


def _infer_filename(file_path_or_buffer, filename_hint: str = "") -> str:
    if filename_hint:
        return filename_hint
    if isinstance(file_path_or_buffer, str):
        return file_path_or_buffer
    return getattr(file_path_or_buffer, "name", "") or ""


def _read_tabular_file(file_path_or_buffer, alias_map: dict, required_keys: list, filename_hint: str = "") -> tuple:
    """
    Dispatch to the right reader (.docx table extraction or Excel with
    smart header detection) based on the file's name/extension. Returns
    (DataFrame, list_of_warnings).
    """
    name = _infer_filename(file_path_or_buffer, filename_hint).lower()
    warnings = []
    if name.endswith(".docx"):
        df = _read_docx_table(file_path_or_buffer, alias_map, required_keys)
        warnings.append(
            "Data was extracted from a Word (.docx) table — please spot-check merged cells "
            "and any values that ran together before relying on the allocation."
        )
    else:
        df = _read_excel_with_smart_header(file_path_or_buffer, alias_map, required_keys)
    return df, warnings


@dataclass
class LoadResult:
    df: pd.DataFrame
    warnings: list


def load_previous_workload(file_path_or_buffer, default_program: str = "", filename_hint: str = "") -> LoadResult:
    """
    Load previous semester faculty workload from Excel (.xlsx) or Word (.docx)
    into a normalized DataFrame: faculty_name, course_code, course_name,
    theory_lab, semester, credit_hours, contact_hours, program, is_lab,
    core_course. Merged-cell/repeated faculty names are forward-filled, and
    text is cleaned of stray newlines from wrapped cells.
    """
    raw, warnings = _read_tabular_file(
        file_path_or_buffer, WORKLOAD_COLUMN_ALIASES, ["faculty_name", "course_name"], filename_hint
    )
    raw = raw.dropna(how="all")

    colmap = _match_columns(raw.columns, WORKLOAD_COLUMN_ALIASES)
    missing = [c for c in ["faculty_name", "course_name"] if c not in colmap]
    if missing:
        raise ValueError(
            f"Could not find required column(s) {missing} in the workload file. "
            f"Found columns: {list(raw.columns)}"
        )

    df = pd.DataFrame()
    for canonical, orig in colmap.items():
        df[canonical] = raw[orig]

    df["faculty_name"] = df["faculty_name"].ffill()
    df = df[~(df["course_name"].isna() & df.get("course_code", pd.Series([None] * len(df))).isna())]
    df = df[df["course_name"].notna()]

    for col in ["course_code", "theory_lab", "semester", "credit_hours", "contact_hours", "program"]:
        if col not in df.columns:
            df[col] = None

    df["program"] = df["program"].fillna(default_program or "")
    df["credit_hours"] = pd.to_numeric(df["credit_hours"], errors="coerce")
    if df["credit_hours"].isna().any():
        warnings.append("Some rows had a non-numeric or missing credit-hour value; they were treated as 0 for totals.")
    df["credit_hours"] = df["credit_hours"].fillna(0)

    df["is_lab"] = df["theory_lab"].apply(is_lab)
    df["faculty_name"] = df["faculty_name"].apply(_clean_text)
    df["course_name"] = df["course_name"].apply(_clean_text)
    df["course_code"] = df["course_code"].apply(_clean_text)
    df["core_course"] = df["course_name"].apply(is_core_course)

    df = df.reset_index(drop=True)
    if df.empty:
        raise ValueError(
            "The workload file was read, but no usable rows remained after cleaning. "
            "Check that the file has a header row (Faculty Name, Course Code, Course Name, ...) "
            "followed directly by data rows, with no extra title rows in between that couldn't be skipped."
        )
    return LoadResult(df=df, warnings=warnings)


def load_scheme_of_studies(file_path_or_buffer, program: str, filename_hint: str = "") -> LoadResult:
    """Load a Scheme of Studies file (Excel, Word, or PDF) for one program."""
    name = _infer_filename(file_path_or_buffer, filename_hint).lower()
    is_pdf = name.endswith(".pdf")

    if is_pdf:
        df = _extract_scheme_from_pdf(file_path_or_buffer)
        warnings = [
            "Scheme was extracted from a PDF using table parsing — "
            "please spot-check course codes and credit hours before relying on the allocation."
        ]
    else:
        df, warnings = _read_tabular_file(file_path_or_buffer, SCHEME_COLUMN_ALIASES, ["course_name"], filename_hint)

    df = df.dropna(how="all")
    colmap = _match_columns(df.columns, SCHEME_COLUMN_ALIASES)
    missing = [c for c in ["course_name"] if c not in colmap]
    if missing:
        raise ValueError(
            f"Could not find required column(s) {missing} in the {program} scheme file. "
            f"Found columns: {list(df.columns)}"
        )

    out = pd.DataFrame()
    for canonical, orig in colmap.items():
        out[canonical] = df[orig]

    out = out[out["course_name"].notna()].reset_index(drop=True)
    for col in ["course_code", "credit_hours", "theory_lab", "semester"]:
        if col not in out.columns:
            out[col] = None

    out["program"] = program
    out["credit_hours"] = pd.to_numeric(out["credit_hours"], errors="coerce").fillna(3)
    out["is_lab"] = out["theory_lab"].apply(is_lab)
    out["course_name"] = out["course_name"].apply(_clean_text)
    out["course_code"] = out["course_code"].apply(_clean_text)
    out["core_course"] = out["course_name"].apply(is_core_course)

    if out.empty:
        raise ValueError(
            f"The {program} scheme file was read, but no usable course rows remained after cleaning. "
            "Check that the file has a header row (Course Code, Course Name, Credit Hours, ...) "
            "followed directly by data rows, with no extra title rows in between that couldn't be skipped."
        )
    return LoadResult(df=out, warnings=warnings)


def _extract_scheme_from_pdf(file_path_or_buffer) -> pd.DataFrame:
    import pdfplumber

    rows = []
    header = None
    with pdfplumber.open(file_path_or_buffer) as pdf:
        for page in pdf.pages:
            for table in page.extract_tables() or []:
                if not table:
                    continue
                start = 0
                if header is None:
                    header = table[0]
                    start = 1
                for row in table[start:]:
                    rows.append(row)

    if header is None or not rows:
        raise ValueError("No tables could be extracted from the scheme PDF. Try exporting it as Excel instead.")

    return pd.DataFrame(rows, columns=header)


def build_faculty_history(workload_df: pd.DataFrame) -> dict:
    """Faculty Member -> prior courses, credit-hour totals, core courses taught, theory/lab counts."""
    history = {}
    for faculty, group in workload_df.groupby("faculty_name"):
        history[faculty] = {
            "courses": group["course_name"].tolist(),
            "course_codes": group.get("course_code", pd.Series(dtype=object)).tolist(),
            "credit_hours_total": float(group["credit_hours"].sum()),
            "core_courses_taught": [c for c in group["core_course"].tolist() if isinstance(c, str) and c],
            "theory_count": int((~group["is_lab"]).sum()),
            "lab_count": int(group["is_lab"].sum()),
        }
    return history


# =============================================================================
# SECTION 2 — RECOMMENDATION MODEL (content-based suitability scoring)
# =============================================================================
# Why content-based + rules, not a trained ML model: a single semester's
# workload gives one labeled example per faculty member — not enough to
# train a classifier/ranker that generalizes. Score = TF-IDF/cosine
# similarity between a faculty member's teaching-history text and each
# candidate course, blended with bonuses/penalties for the constraints
# in the brief (Computing Core Course coverage, repeat-avoidance).

@dataclass
class SuitabilityResult:
    scores: pd.DataFrame
    explanations: dict = field(default_factory=dict)


def _faculty_profile_text(history_entry: dict) -> str:
    return " ".join(history_entry.get("courses", [])) or "no prior courses on record"


def _course_text(row: pd.Series) -> str:
    parts = [str(row.get("course_name", "")), str(row.get("program", "") or "")]
    core = row.get("core_course")
    if isinstance(core, str) and core:
        parts.append(core)
    return " ".join(parts)


def compute_suitability_scores(
    scheme_df: pd.DataFrame,
    faculty_history: dict,
    repeat_penalty: float = 0.5,
    core_bonus: float = 0.25,
    core_repeat_penalty: float = 0.3,
) -> SuitabilityResult:
    faculty_names = list(faculty_history.keys())
    if not faculty_names:
        raise ValueError("No faculty history available to score against.")
    if scheme_df.empty:
        raise ValueError(
            "No candidate courses were found in the Scheme of Studies file(s) after cleaning. "
            "Check the uploaded scheme file(s) have a recognizable Course Name column."
        )

    course_texts = scheme_df.apply(_course_text, axis=1).tolist()
    faculty_texts = [_faculty_profile_text(faculty_history[f]) for f in faculty_names]

    vectorizer = TfidfVectorizer(stop_words="english")
    tfidf = vectorizer.fit_transform(course_texts + faculty_texts)
    course_vecs = tfidf[: len(course_texts)]
    faculty_vecs = tfidf[len(course_texts):]

    sim_matrix = cosine_similarity(faculty_vecs, course_vecs)

    rows = []
    explanations = {}

    for fi, faculty in enumerate(faculty_names):
        hist = faculty_history[faculty]
        prev_courses_norm = {c.strip().lower() for c in hist.get("courses", [])}

        for ci, (_, course_row) in enumerate(scheme_df.iterrows()):
            base_sim = float(sim_matrix[fi, ci])
            score = base_sim
            reasons = [f"content similarity to prior teaching = {base_sim:.2f}"]

            course_name = str(course_row["course_name"]).strip()
            core = course_row.get("core_course")

            already_taught = course_name.lower() in prev_courses_norm
            if already_taught:
                score -= repeat_penalty
                reasons.append("penalized: identical course taught last semester (Rule B)")

            if isinstance(core, str) and core:
                score += core_bonus
                reasons.append(f"bonus: is a Computing Core Course ({core}) (Rule C)")
                if core in hist.get("core_courses_taught", []):
                    score -= core_repeat_penalty
                    reasons.append("penalized: faculty already taught this exact core course last semester")
                elif hist.get("core_courses_taught"):
                    reasons.append("note: faculty taught a different core course last semester — this diversifies core coverage")

            score = max(score, 0.0)
            rows.append({
                "faculty_name": faculty,
                "course_name": course_name,
                "course_code": course_row.get("course_code"),
                "is_lab": bool(course_row.get("is_lab", False)),
                "credit_hours": float(course_row.get("credit_hours", 3) or 3),
                "core_course": core,
                "program": course_row.get("program"),
                "already_taught_previously": already_taught,
                "score": round(score, 4),
            })
            explanations[(faculty, str(course_row.get("course_code")))] = "; ".join(reasons)

    return SuitabilityResult(scores=pd.DataFrame(rows), explanations=explanations)


# =============================================================================
# SECTION 3 — OPTIMIZATION (constraint-based final allocation, PuLP/CBC)
# =============================================================================
# Decision variable x[f, c] = 1 if faculty f is assigned course c.
# Enforces: <=9 credit hours/faculty, <=3 theory + <=1 lab/faculty,
# <=1 previously-taught course repeated/faculty, <=2 faculty per course.
# Objective: maximize total suitability score + a soft per-faculty bonus
# for having at least one Computing Core Course (kept soft so the solver
# stays feasible when core seats are scarce).

@dataclass
class AllocationResult:
    assignments: pd.DataFrame
    unallocated_courses: pd.DataFrame
    faculty_summary: pd.DataFrame
    solver_status: str


def optimize_allocation(
    scores_df: pd.DataFrame,
    explanations: dict,
    max_credit_hours: int = MAX_CREDIT_HOURS,
    max_theory: int = MAX_THEORY_COURSES,
    max_lab: int = MAX_LAB_COURSES,
    max_faculty_per_course: int = MAX_FACULTY_PER_COURSE,
    core_course_soft_bonus: float = 2.0,
) -> AllocationResult:
    faculty_list = scores_df["faculty_name"].unique().tolist()

    scores_df = scores_df.copy()
    scores_df["course_key"] = scores_df.apply(lambda r: f"{r['course_code']}||{r['course_name']}", axis=1)
    course_keys = scores_df[["course_code", "course_name", "course_key"]].drop_duplicates()

    prob = pulp.LpProblem("faculty_workload_allocation", pulp.LpMaximize)

    x = {
        (row["faculty_name"], row["course_key"]): pulp.LpVariable(
            f"x_{row['faculty_name']}_{row['course_key']}".replace(" ", "_"), cat="Binary"
        )
        for _, row in scores_df.iterrows()
    }
    has_core = {f: pulp.LpVariable(f"has_core_{f}".replace(" ", "_"), cat="Binary") for f in faculty_list}

    prob += (
        pulp.lpSum(x[(r["faculty_name"], r["course_key"])] * r["score"] for _, r in scores_df.iterrows())
        + pulp.lpSum(has_core[f] * core_course_soft_bonus for f in faculty_list)
    )

    for f in faculty_list:
        f_rows = scores_df[scores_df["faculty_name"] == f]

        prob += pulp.lpSum(x[(f, r["course_key"])] * r["credit_hours"] for _, r in f_rows.iterrows()) <= max_credit_hours

        theory_rows = f_rows[~f_rows["is_lab"]]
        prob += pulp.lpSum(x[(f, r["course_key"])] for _, r in theory_rows.iterrows()) <= max_theory

        lab_rows = f_rows[f_rows["is_lab"]]
        prob += pulp.lpSum(x[(f, r["course_key"])] for _, r in lab_rows.iterrows()) <= max_lab

        repeat_rows = f_rows[f_rows["already_taught_previously"]]
        prob += pulp.lpSum(x[(f, r["course_key"])] for _, r in repeat_rows.iterrows()) <= MAX_PREVIOUSLY_TAUGHT

        core_rows = f_rows[f_rows["core_course"].notna()]
        if not core_rows.empty:
            prob += has_core[f] <= pulp.lpSum(x[(f, r["course_key"])] for _, r in core_rows.iterrows())
        else:
            prob += has_core[f] == 0

    for _, crow in course_keys.iterrows():
        ck = crow["course_key"]
        assigned_vars = [x[(f, ck)] for f in faculty_list if (f, ck) in x]
        if assigned_vars:
            prob += pulp.lpSum(assigned_vars) <= max_faculty_per_course

    prob.solve(pulp.PULP_CBC_CMD(msg=False))
    status = pulp.LpStatus[prob.status]

    assigned_rows = []
    for _, row in scores_df.iterrows():
        var = x[(row["faculty_name"], row["course_key"])]
        if var.value() and var.value() > 0.5:
            reason = explanations.get((row["faculty_name"], str(row["course_code"])), "")
            assigned_rows.append({
                "faculty_name": row["faculty_name"],
                "course_code": row["course_code"],
                "course_name": row["course_name"],
                "credit_hours": row["credit_hours"],
                "is_lab": row["is_lab"],
                "core_course": row["core_course"],
                "already_taught_previously": row["already_taught_previously"],
                "score": row["score"],
                "reason": reason,
            })

    assignments = pd.DataFrame(assigned_rows)
    if not assignments.empty:
        assigned_keys = set(assignments.apply(lambda r: f"{r['course_code']}||{r['course_name']}", axis=1))
    else:
        assigned_keys = set()

    unallocated = course_keys[~course_keys["course_key"].isin(assigned_keys)].copy()
    unallocated["reason"] = "No faculty member could take this course without breaking a workload constraint (credit cap, theory/lab cap, repeat cap, or duplicate cap)."
    unallocated = unallocated.drop(columns=["course_key"])

    faculty_summary_rows = []
    for f in faculty_list:
        fa = assignments[assignments["faculty_name"] == f] if not assignments.empty else pd.DataFrame()
        faculty_summary_rows.append({
            "faculty_name": f,
            "num_courses": len(fa),
            "total_credit_hours": float(fa["credit_hours"].sum()) if not fa.empty else 0.0,
            "theory_courses": int((~fa["is_lab"]).sum()) if not fa.empty else 0,
            "lab_courses": int(fa["is_lab"].sum()) if not fa.empty else 0,
            "has_core_course": bool(fa["core_course"].notna().any()) if not fa.empty else False,
            "repeated_course": bool(fa["already_taught_previously"].any()) if not fa.empty else False,
        })

    return AllocationResult(
        assignments=assignments,
        unallocated_courses=unallocated,
        faculty_summary=pd.DataFrame(faculty_summary_rows),
        solver_status=status,
    )


# =============================================================================
# SECTION 4 — PIPELINE ORCHESTRATION
# =============================================================================

def _derive_scheme_from_workload(workload_df: pd.DataFrame) -> pd.DataFrame:
    """
    Build a candidate course pool directly from the previous semester's
    workload when no separate Scheme of Studies file is supplied — used as
    a fallback so the tool still produces a recommendation from workload
    data alone, at the cost of only knowing about courses that were taught
    before (it can't suggest brand-new courses it has never seen).
    """
    keep_cols = ["course_code", "course_name", "credit_hours", "theory_lab", "semester", "is_lab", "core_course"]
    available = [c for c in keep_cols if c in workload_df.columns]
    df = workload_df[available].drop_duplicates(subset=["course_code", "course_name"]).reset_index(drop=True)
    df["program"] = ""
    return df


@dataclass
class PipelineResult:
    workload_df: pd.DataFrame
    scheme_df: pd.DataFrame
    faculty_history: dict
    scores_df: pd.DataFrame
    assignments: pd.DataFrame
    unallocated_courses: pd.DataFrame
    faculty_summary: pd.DataFrame
    solver_status: str
    warnings: list


def run_pipeline(workload_file, scheme_files: dict, workload_filename: str = "") -> PipelineResult:
    """
    workload_file: path or file-like object for the previous semester workload
                   (.xlsx or .docx)
    scheme_files: {"BSAI": (file_obj, filename), "BSCS": (...), "BSSE": (...)}
                  programs with no file supplied should be omitted or set to (None, "")
    """
    warnings = []

    wl_result = load_previous_workload(workload_file, filename_hint=workload_filename)
    workload_df = wl_result.df
    warnings += wl_result.warnings

    scheme_frames = []
    for program, (file_obj, filename) in scheme_files.items():
        if file_obj is None:
            continue
        sch_result = load_scheme_of_studies(file_obj, program=program, filename_hint=filename)
        scheme_frames.append(sch_result.df)
        warnings += [f"[{program}] {w}" for w in sch_result.warnings]

    if not scheme_frames:
        # No separate Scheme of Studies was supplied — fall back to using the
        # previous semester's own course list as the candidate pool for the
        # upcoming semester, since that's the only course data available.
        scheme_df = _derive_scheme_from_workload(workload_df)
        warnings.append(
            "No separate Scheme of Studies was uploaded, so the previous semester's own course "
            "list was used as the candidate pool for the upcoming semester. Upload a real Scheme "
            "of Studies for Fall 2026 for a more accurate result."
        )
    else:
        scheme_df = pd.concat(scheme_frames, ignore_index=True)

    faculty_history = build_faculty_history(workload_df)
    suitability = compute_suitability_scores(scheme_df, faculty_history)
    allocation = optimize_allocation(suitability.scores, suitability.explanations)

    return PipelineResult(
        workload_df=workload_df,
        scheme_df=scheme_df,
        faculty_history=faculty_history,
        scores_df=suitability.scores,
        assignments=allocation.assignments,
        unallocated_courses=allocation.unallocated_courses,
        faculty_summary=allocation.faculty_summary,
        solver_status=allocation.solver_status,
        warnings=warnings,
    )


def faculty_teaching_map(result: PipelineResult) -> pd.DataFrame:
    """Faculty -> Previous Semester Courses -> Recommended Upcoming Courses mapping."""
    rows = []
    for faculty, hist in result.faculty_history.items():
        recommended = result.assignments[result.assignments["faculty_name"] == faculty] \
            if not result.assignments.empty else pd.DataFrame()
        rows.append({
            "faculty_name": faculty,
            "previous_semester_courses": ", ".join(hist["courses"]) if hist["courses"] else "—",
            "recommended_upcoming_courses": ", ".join(recommended["course_name"].tolist()) if not recommended.empty else "None allocated",
            "repeated_from_previous": ", ".join(
                recommended[recommended["already_taught_previously"]]["course_name"].tolist()
            ) if not recommended.empty else "—",
            "core_course_assigned": ", ".join(
                [c for c in recommended["core_course"].tolist() if isinstance(c, str) and c]
            ) if not recommended.empty else "—",
        })
    return pd.DataFrame(rows)


# =============================================================================
# SECTION 5 — PDF REPORT GENERATION
# =============================================================================

def _pdf_styles():
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="H1c", parent=styles["Heading1"], spaceAfter=10))
    styles.add(ParagraphStyle(name="H2c", parent=styles["Heading2"], spaceBefore=14, spaceAfter=8))
    styles.add(ParagraphStyle(name="Small", parent=styles["Normal"], fontSize=8, leading=10))
    return styles


def _pdf_table(data, col_widths=None, header_bg=colors.HexColor("#2F5496")):
    t = Table(data, colWidths=col_widths, repeatRows=1)
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), header_bg),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F2F2F2")]),
    ]))
    return t


def generate_pdf_report(result: PipelineResult, semester_label: str = "Upcoming Semester") -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=letter,
        leftMargin=0.5 * inch, rightMargin=0.5 * inch,
        topMargin=0.6 * inch, bottomMargin=0.5 * inch,
    )
    styles = _pdf_styles()
    story = []

    story.append(Paragraph(f"Faculty Workload Allocation Report — {semester_label}", styles["H1c"]))
    story.append(Paragraph(f"Solver status: {result.solver_status}", styles["Normal"]))
    story.append(Spacer(1, 12))

    story.append(Paragraph("Faculty Workload Summary", styles["H2c"]))
    header = ["Faculty Member", "Courses", "Credit Hrs", "Theory", "Lab", "Core Course"]
    rows = [header]
    for _, r in result.faculty_summary.iterrows():
        rows.append([r["faculty_name"], str(r["num_courses"]), str(r["total_credit_hours"]),
                     str(r["theory_courses"]), str(r["lab_courses"]), "Yes" if r["has_core_course"] else "No"])
    story.append(_pdf_table(rows, col_widths=[1.8 * inch, 0.7 * inch, 0.8 * inch, 0.6 * inch, 0.5 * inch, 0.9 * inch]))
    story.append(Spacer(1, 12))

    story.append(Paragraph("Faculty Teaching History", styles["H2c"]))
    tmap = faculty_teaching_map(result)
    for _, r in tmap.iterrows():
        story.append(Paragraph(f"<b>{r['faculty_name']}</b>", styles["Normal"]))
        story.append(Paragraph(f"Previous semester: {r['previous_semester_courses']}", styles["Small"]))
        story.append(Paragraph(f"Recommended for upcoming semester: {r['recommended_upcoming_courses']}", styles["Small"]))
        story.append(Paragraph(f"Repeated from previous semester: {r['repeated_from_previous']}", styles["Small"]))
        story.append(Paragraph(f"Computing Core Course assigned: {r['core_course_assigned']}", styles["Small"]))
        story.append(Spacer(1, 8))

    story.append(PageBreak())

    story.append(Paragraph("Recommendation Explanation", styles["H2c"]))
    if not result.assignments.empty:
        for _, r in result.assignments.iterrows():
            story.append(Paragraph(f"<b>{r['faculty_name']} → {r['course_name']}</b> ({r['course_code']}): {r['reason']}", styles["Small"]))
            story.append(Spacer(1, 4))
    else:
        story.append(Paragraph("No allocations were made.", styles["Normal"]))

    story.append(Spacer(1, 12))
    story.append(Paragraph("Unallocated Courses", styles["H2c"]))
    if result.unallocated_courses.empty:
        story.append(Paragraph("All scheme courses were successfully allocated.", styles["Normal"]))
    else:
        header = ["Course Code", "Course Name", "Reason"]
        rows = [header]
        for _, r in result.unallocated_courses.iterrows():
            rows.append([str(r.get("course_code", "")), r["course_name"], r["reason"]])
        story.append(_pdf_table(rows, col_widths=[1.0 * inch, 2.2 * inch, 3.1 * inch]))

    if result.warnings:
        story.append(Spacer(1, 12))
        story.append(Paragraph("Data Warnings", styles["H2c"]))
        for w in result.warnings:
            story.append(Paragraph(f"• {w}", styles["Small"]))

    doc.build(story)
    return buf.getvalue()


# =============================================================================
# SECTION 6 — STREAMLIT FRONTEND
# =============================================================================

# Iqra University (Chak Shahzad Campus, Islamabad) logo, embedded as base64 so
# this file stays fully self-contained (no separate image asset to ship).
IU_LOGO_B64 = "iVBORw0KGgoAAAANSUhEUgAAAXgAAABoCAYAAAD7POuJAAAAAXNSR0IArs4c6QAAAARnQU1BAACxjwv8YQUAAAAJcEhZcwAADsMAAA7DAcdvqGQAAIhPSURBVHhe7b13VFXJtj56/35vjPd+7/e79557Up/u0znnbnObc845tdqGNrRZEUEUCSKCooggZsSAIqCiiCJiRAVRERUFAck5G783v1p7wWa7QVRM5+waYw4U1qpVNWvWV3POmjXrPz7vaIOP2y7AN10X4esuFnpZ9BWpM8lW/fyaZPKMhSz0WknJpy6jFnobyHQMvxH6spMtPmg1H62HuOCNAXg29gtpy2cdFlbR589ANd6TehSImvnOyyTykMRv63wlo//Zch7e+4U01/BTI/6ef/+ojRU+a79Q8YAD9DrGgkLBNtfg4wsSx/NLCmKXakFUPDL67ttCHBt9TPnzdc+XhiYqIB+3s8b7LefXkFELvdlEefyyk8ij0Vi+kQD/tYAAweALaSyJAPFpe+t6EwGJ77Gz7OCrBnh+j4DGtnzSlm1ie2zxbVc7/NRrKZr2d8Ivg5zRcrCL/FyGZgOc8HPvpfiu+2LVbz7PcfhEJhnH5HW0n3xTvJfvk5+fmPD4mUj68Um7BQIaC/CR9Evvm+KL6t+T2sebTGwvF6pvZDz1/5s+87YSwf1r6RflsZnIaYuByzQa4CzEnxZ648gwRo36OOC7bovF+tLkkTj+RgE8Jw3p2252aNrPEW2HuQotR6vB7ICzAsKnUXN5rqU830be47sEUgKnue+9LPqio60CsI/aiIbXwUb64oTeE9Zg/IItsFq+F07rQrFq01Gs2RoBjy1H4eJzCDYrAjHZ1g8Dp6xT7SdoUJunJqWBoPlvNTTxO5zcLWXhIf9bD+UiVE/e10Icy5962StAJD/Yp4/aWikin76Q373qRex5iXOECzVlcuAUb7QZulx+t9Dss28bEQxI38t8GTN3E5auOQC3jeFY4XsEruvDLPSGktuGI3DfdAQTrLcqpZFKFRVFzrc3CuB1jbFJX0dMtfPHpj2n4B98FiFHL+Hg8csIOhKLwMMx2HvoIvYeNiLD//n3AxGXcVBo18EL6v3FHiFqEhJACC7mvtsQRJ4RiP8pptL7LefJaroUI2b5YrnPYeyR9p2NvY3ktFzk5JegsLgcxSWV2s/SCpSUVSKvoBR3Mwtw5cZd7D8Wp4B/ks02BfYftKapPFdpwvq3TL//okT+0MogjwZM9oKTVyh2HjiPkHDhvfCTfA0Mq4P38tzhE1dx/Ox1RAjxZ+S5G4g4c129u0vq2rj7FDy3RcBx7QFMX7oDA2QxayLgT5l7t8Vc5Zoy17Y3iTi2XLBmLN2J8JPXsGB5oPo9rRxOJtPn3ybi3GM/uMB7+0fiZnIWsnKLkZlTZKE3mLLzihWubNx9En0meuL91lZqLN84gGdjCDI0NUbP2SggcxBuvmECIDGIv5mOsvJ7KBFALBJgLCp5kioq7yMlPR+hx68YgOQgptvvUBqoAjDpn7nvviixbvKNRE11+Iz1WL7+MA5HXVWTpLS8EnpJzypEdFwyjgg47D92WYHimZjb6rn79x+qZ9iPO+l5OBF9A+tkov1mtUVZMPo3yCPdPdBQxD4o/guPeozzwGzH3aIVhGPbvjM4J4uTtiiZ5zupsKhcLVDXb2fiRlIWEu9k487dPAUQ7D/HjotYWkY+Em5l4IzUGSwL96rNXMj80EoEkADP/r2pLg+26x/N56Dnb6uxIyQaZWX3sHN/NDqMXKHtnQg4NvS4vEqiJULtj/OPC3J5xT0lj5bydhQquPQAvNdy/psJ8CQCDRv2fY/F+LHnEvwkNHiaDzbvPYPKew8MXam9nIi+ibHzN6P5AGcxNZfghx5LXuqk46RXPuq21qCfcrbjLhw9lSCgV2FokVb4/1MXbinNeNBUumFc0FgmEttJwJjlsFsG6JJaAB49emx4C3go/+aCsGT1frQfsQKfyrfouyafzLWnIejbbouF/0vwoyxWBC/ndaHK+jBul2nhonv0dIKY82HK+lgjC+wG0SgodFycC4rL8PDhI8PT1eXe/QeIiU+Bi1g6HUa6qf2WzwRkXmb/nodo8n5mcC/9sWSHWshYouOS5P/+2uLbboHZd98W4vynbBHgN4v1S63QUt6ewkW57+9r8c83GeBJnNwETmqTn7ZbqMzigVO9cOpiYp1axbXEdNi6BymXBicbwYIdVBtHL4G0DVFr5T4hSFPbpsl0TzTxx0ZYmCG/8ws6h25jPVR76JvnoqBIJhT7SL7THbNwxT5cv5WJBw+qwfCBAGN6dqG2QsviQA2ePDHXpoYgXZsn//7ebDbaDV8Br+2RTyxaenkknaVraZYsbuQF5Yjvsy6CRbexq7DQbR+On7teZaUYF76fX1QGv+Cz6D/FS7m5XuW+w9OI84G8oJZOd5/PjhOy8Grjkysg6Bd0Fj+IIsIxfFlW4qugGgC/97QF4Osotas6r6+8NQBP0gGeIPiXxjNFe3VV/ne6CmordGlMtN6KvzaZpU02PYrGTP0NQQRqatT9hKlbAk8rEDYuFAK6J2jG9520VgEEfc0fGsx5msT8+bFo/wyT5N/oinFccwC37mRrlRiV3IJStRcxZt4mFaf8YWsr1b+GtlB0gCf//9xoplhS9li0Mhh5haWGltQs1Oxjr6Viku02/J8fpuGd5nNUX94T0sI/rZRfd7yMzemLt5S7xlxh/wie7WVB4cLJxe9lWl/1JpkP3Awmvycv2o6T528aWqyVmPg7SgH5srONGmOzdbwFZArwXLyMS0lppVrIz11KUntK5y8nI+ZqCuIS0tTvSZevvxhdMfy8dC0NF6/cUdYr3Xn8Jq1IKnhUCOh3viQyd+pCorKiLko7+I56/0XaUdWPNGVZso/8NpVLficxObtKSamsvI+ExAz1d8p1tPxUPJH3+L6q53nbYmgHif0kL87J99lf/p+KpKlF/FYBvE6MqnlPQKLTaHfsEKAsEE2vtkK/Nv3V7NDLjmzQAajjKHds3XfmicnA8uDBQ5yOuaWAjYDHyc/QTUZhGNdFfjNyiMD/vrSdrp5dtfT1sQg3NzuHz/RVCwMXiJeh6eoyQJDmhrH1ikDlQzdXCPAURPrt2X5tB19ztXGhpkz9rclsxa+la/aLpZVhePPJQv88rTC+z/dM2/U6iIoCAZ7RWL67opCRXWRorVayc4vg7HUQzQY4KgvGXB1vA9UF8JQ77qms3hoBa9d9mL9sj7I2rV0DFfHfVALsVoUYiP9+FtLe4dhbS10LDHWS5sm3+DtasAQ2WrQXBPBWbgxXLlEbsQ7ZJrbD5oXaEYJF8lNvg94vG7cgzHMOgP3q/dgtIMqFjoXzYXPAKdjJ96xc9sJWnqviCd+Teljf87RDtUXqrarP0JY5TruxalO4WvQq791X7dDLWwnwBAsG8XcevVJFddQN8AkYZ7VVgQxNanP1NQQRtAjYDCOkkNXmuqC1wQFj3DsnPt8zV58xMV6c7hdG0DAKxVwpK7+PfSLsyicvE5Lg87LGjP3kxvF8lz1KwzZXdICni4YaOwHR2GrSwZoWT/dxqxAUHmt488lSKtp9+Olrqm+MvecCZtye10FcmDk3+k7yVNqaaaFFwqihfr97ijIyV/XfXD1vOtUF8NSar95MVxZM1zGr0EHGp+MoNyUbHFtdgeH7jMTiYv5MJONMpYlzl3OF+zDNBzrJvHeXb7mh668r1d5OkmjxVJyOnU5QeyFU/OgCbNbPUdXzXos5qj3897O2g/3nXhoB8YPW8/GDWK4dpI99ZNwZeDBq1gY133UM4kLPvbGh033Qa/xq9J3oqfaRuO9Hq/X9VvPUfGa95r5XG+nt5nwhX8mbpv0d0VX62UVwcLr0m4EkZSbu6n8LgP/tJQI8+aJrc193sZNVdR9up+QYvl6zcEOY5hpdM9pBHy3EsS5i/bQ8OCD0x6/1i6h1Y5PaFFdyhhnSZ12fxeN5SAd4alHPD/B2GngID37sYS/9Oq6AXDDDbKEpPs5qs2bVtH79Lo93BXAa9XXAUs8DSM3IN7SyulC7zS0oUWcceNBEs9Q0K8ZcfW8qqTESYDGvwUPNvdMxtxEWFY/QyKs4EBGn9h+oWXYRAP5AgOUfArA830AcUYt6PYj7ZQRD8pmyQ8BkfPeug+fVdw6duCrfvKo26xlFR35Tk6cWT2WAkWgBoRewxCNEATHnP+W23t83EN+jcvWzyPvMpbuwftdJHBAg1cJ+byhrPOF2htpjY2G0G11CDJ44rkKCE7A/4rK8F6XCaAn0eth0ffihP0PSFswFsnCsgcOaA6p/h4XvXNjorknPKvjXcNE8C8CHnxKAX/ByAZ5aBhk/5A8fFd9dW6EAMhKBMf1keH0BmP3VtRm6PDKya0bV6IWDG5+YjlGzNyhfOReQl+Gq0QH+RTR4EgWNY0LNxNkrVPlQOVHNlfzCUmV2K97J2BvX86qJvKU89RbtjFFCpWXVIa+mZXvwOfSe4Kl4wLlE3725Ot9UqgvgzRVq9eUV91VEEd2UDNGjG+tjkV3WRxmoD1E2aK0xaovAyHMgd7MK6xU1p5f7otVT2WIwQ4/fVqt5pOaEme+ZJXme85pAuNA1EOfjksUCl8XEUH99C0Waljstvfkue9VJdbpeWb/Z7xqRGgeZw5z7/HffiWuxc/95pfAYB13UVv71AZ4a/EsEeA7SJ8IfplKgdp1iRpvTCyN66Kck6JLhPP5trk5zxEEhsI2Y4as2VbiZY648evxItYOmMp8nr8zV9yLUUABP3tH0pHzRl8lw0NoAnmPs7B2qrJj3WtA6qVnXqyLyk+3lyU5aMPmFtcsey6XrabCWhYnAou2NvByr6mXRswK8cWG46zbR5nuMXw3umZmOf11ENwZ5PWiqt9ospY/9eQs3Yd03HlEnsN9tPsfs90yJ48S9MWrRUxZtV5uYtVnO9S3sw8WrdzBh4VZVLxWFeskD8ULmMgMNPDYdfaZFzgLwL0gE9y86LhQNwUOZa9Rgais8vDRypq9qB90MZLa5Os0RBYHAyrBKaoX6po5pIUCeu5SMaYv91SAyfPJZJlZ9qMEBXrSk+gD8Mu9DGsDL2L8egLdTk5KbzPTB+odEG1pXeymXhXi3mNItBjlX+VPfpmRqLwLwnAvXbmUoWaQmrGSgPoAmRJlh9NgSj/0qXPZFCiNceMJ40LR1eKf57Hq1gW1lNBzPflDueHCvIQr74uh5QJ3noVbO75j7vjERkLlAjhDsOHIyXu031LdYAP4FiYLIkMHZTruVD7yuwk2QLqPd1SSnRvesPCWwth66HCs3cVPHvMARH/NEq/TYfBTfd1ustEaCkrn6npcazEVjAHj64Z3q4aJhJAM3ll4XwLO95CcjlWzdg1Xomn6SmkBeW6EGyqgp1vEm7B88C70owHMfgqlBOF7kHcHF3HeMiQDMTVVu3PruOqn82i9SCPCxMgaMpiNQaguN+W+TuAAr3BO5pDvFZ2f1GYcXLXSj8kT9L+oUupY4sC4cIC84f7lBy83s2HixZiwAX11eJsCT+Qp0h7kqIagtJpyFQuovmjdzlrAdzxOyyW/xfe7S1/Utln1hseg0yl1pCaSGHL+GAHi2RwMPa+Xu4CnXkjJulhkqMCkp6XmYKKYt/bkUTuP2vAripFcauIwdrYhA4W9ufqkKTaPLrK7FnXsmPrtOqCRrbHt9tdg3gbQxej6A52JNzZeJ85gxlUoNweppFowG8POUlcT5TZ/+ixT64q9cT1OuFloS7FNdY0DZ1BdyRsBsDjytXJ8NVRhWy2g71RZRcOoD8HyWexEMJ34Wd5UF4F+A6KNjvRTEQyeuqARhtRWaZjzVSgazHXzXXJ110bvN5yprgb7fp50o5C7+mLmb1eSkcDyLO+hp9LwAb1wH4/4pVx+0no/uYz1USobaSnnlPXXilecLuCAQaI3rehVE/lH75oSbvMgPN5IykZZZgDnOASq2mVEVtRX6omOvpah0FAz1Yx/MfeNNpBcF+FJZtN03HEGTvg5qvDWQMf8tnXSAJzDRvVWXdVSfQo03/tZd5SrS21BfgO8oAL8l8EytluXzlE0BJ1XIrwq1le/UB+DZ7lkOu9TmtQXgjcrLAniN8Vo6ASZBi0tIxb06Nj+oxblvDBcma2Z+fUxVU2KfqcEy1pcJu+oqPDXHHXv64BvaLVBfgGcbqHX8Q55nnwnyJI4DBYwHnb6S9jmvO1SVx8Vc4d9o5jPckPyua3K+LNInBf3CnDSUOe65MAU1tTxqZXUV5t1hH5hCmQseXQRv0pyqjV4c4Cs1gJf3dXCtH8DPVSk4mHzuRV00CuAT058D4Begy5iV2B5yzlBTwxRG0vHMQH0AXrkxpb20fjjfktPynsldZAH45yQKiG6y0zfGjIgEtdpKipjwLt6H5Xkt1cLzADwHiZrvxIXblPZYV2H2RiYE4/O65mmuzueh+gI8j6tzUjGtBMPC6NOsuiRE/s3EavSLMnystknMKBUmKKMmxXfUJqX0yVy7XhoJGHMznfIzbMZ6JKXmqgV7nf9x9Xvyl7y4m1VQqwzw9yFH4zBkuo8Cr+cZ/9dBnP8vAvB0uxHgOdbPCvCMfWc8+7NEjZgrBPirAvBTnwHgleUrzzK8cvfBC4aaGqY8iwavAzznDM8W3M00HyJdW7EA/HOSDvC8yGOmmE4FTwmXu52aozYSCe4K4OsQsNqIg8JIDrpebtVymEovBCGm3f2hx+LXBvB00fDYOA91sL/fdrfDj72WiBbspHJUO609qMLG6MIwV/KLSuG7M0olbVNuHpFFgq25Nr1M+qqTlp6BF564eB9SgMEcR8z/w5DNvzSepfLYhxyLq9VfTH6kpudhyeoQ1Q9tkW/Yze+XQRq4vLiL5lkAnrLyrsxvpt7gOYPXBfAftlmg5HTf4RhDTQ1TXgTgTTPLPq1YAP45iQJCIaDLYZbjbhTXkppAL0kNAPDK1SIAN3LWhjpdGizc9PPcdlylVn4dAM/JTTCIPHsDPv4n4C1ETdxfzF0mR+OJXlo9xpo73+H/U9PzEXYiXuXeoN+dAqkOqLwGcCe/v5Q5wD4T0Jl6moCzVnjL6+tombzTbI7aQF20MuipAMiLXmj286QmJ+0rt0aekTRweX0Az4OKrxvguaHekMUC8E+hNwbg2xHgbdXp0tpyz+hFA/iDLw7wXWzVrVBMwFVXoa+O1/8x3vZ1ADyLAmsBcR4nT80oUPsGHKvaJiwTNfEkME+1chFjvh7KHul1gDujPaht60DMa9DYX5V/xdYQkSFjyYnK55hBlBYJAaW2cikhDfOX75XJpr1n7rtvEmng8u8L8DytzIy1DVleyEVjAfia5eW7aGxUjoqnhS0mpQnAr2sIgF+kUhHQx15XoQvHbeMRFYL4ejR4bixqx7N5vSKPjD9tovKZuc4BaiP5fxrNUEL4Ot0YaqKL5UD5pyYXfuqamlwbAk6pEFTygXzlhCGPf+7tIJP3lNKyais8o7DnUIxo/45gfhbT6KI3jTRwsQB8QxYLwD+F3hiAN4D1JFs/JN/NM3s7kV4Yx8144KpN1joErDZiHzhAv1ltfuqhqoTbmSoRFp9XC8NzfK82qg/AUwjjEzMw2ylAtYGX/56NvW34q/nCmGlmleSpYF6DR3BtyIXpeYguGE4uutfIcwIc+/Jdd+FrG46HpumrySpz5XfR7JmzpK5C9xpPJfJcAOs39903hTRwsQB8QxYLwD+F3hSAV2a28IYuE2axqyuci4nGVm46qpjMdjwPwDMfBbMpMg6bURx1FebOILh+LhpiQ4NIfQGe7ogJC7fhf383TWmqvA+3rs1hvkM+8Yo/HgSpFkrz7XjZxIWYvnLeDcs0zWwbD5AxI+CfxcL4W9NZSg5JvFCG0UK0Prz8tJuuaMWYKzwTQeuK9TKE9LXsLdSTNHCxAHxDFgvAP4XeBIAnqYNOLeepvM90QxSX1u6H5wXUPjuiFIMVwD+HaU5gJbgwLcLTDjoRkEbM2qAElYsQv2uuzueh+gI8b6CZ6bgbf282RwFg519XqhO/td3epBfGzzMtAfc3GFb5utwYTFdLTX2ijR+y80qU64VnGXiamK4vbrKSuMHK3/ESFB5E40J2OyVXAZy5QsA6eTFRKQbvy8Sjm+9NBXkNXCwA35DFAvBPoTcF4CkkBLtWg12wZuuxOkG3QoSUyamYS4W+++dpC0+yMtc7D8zUdpMSC8effOEhHAoqqSHHr74Ar59kZUwzZejD1vPRf/JanJDF52kgz5O4Q//wUVr0B62fLTFbQxDHlnnMecE4/eo8pcy+hp2MxwaZoBt2R2GLAN7mvacMdBqb9pxSp5V5wXpmTu0H0Yj71PAdPA8qPpI3XMSedoT/dZAGLhaAb8hiAfin0JsC8CRqed93X4Kpdv4q9ryuwksRuo/zMJj+z5gfhkIvwNpmqCs8thyr9R5aTioKAQ85KeHoUH3ZdUPRswI8Y9gpR/zJuPzp9jtUZE1dhZvWAYcuKoClMHJRNNeWl0Gc/GqMRO6Z2pURS7zQoUh4zvs3E5OzkHgnG7dSTEh+x81vygEX+6dNRF6zOHiat5JNfu9VL2L1IQ1cLADfkMUC8E+hNwng9TzfjKrgDS9MbFRbYez3mHmbFaO58flNl/pPaN1a6Dl+jYqlruuSCR4OGWu1Gf+UPjc0uJOeB+C/kPZTSClLdGV474jE3aecxqXPm7HldIOo3PYvoS+mpIdGcnx4sImWGcGdC1Jw+CUFcrxJh2MQEPok7RHaGnhGnX5kYqi6EmUxla6D5wG1CJM3dYHO6yINXCwA35DFAvBPoTcJ4Ok7JX9oYvNSgdqu62Nh9AQv62X6APpezdVnjiiMrJ994GGbmDpOf/LSXaZE4AaelvPEtsFN/+cCeIMfnROH/WCCtl1POQJOgOCt9Vpue20yvAoQpLXANjJa5uT5RBUdxY1fTkrOh59626tF6ufe5shePdNysAtWijyYu85PLxUyVvsj4tRFFJ91eDOTkGngYgH4hiwWgH8KPRPAi4C8TIAneFIYyKN+k9aqKIvaCg/6bNt3Fk37MeFU/bRr1s/n9JOc3HxkfDknj2mprHyAc7FJ6P+7l7Tn5V1O/SIAz75QWHmZAnPq0KqpKzsec4pTc+bl1pwQvBzkZboyqL2T1/w3E4hl5hYpd8zgP7yVVs+JwfEmb2sjRsZwc3jo9PWyQNUdMkktforddnzbTXP3mbbndZMGLhaAb8hiAfinkA7wneoB8C/90m0hMoxaH4VmjlOAmPPmT5lS6465mqIAmCBCvup1mNarEweEoPHPltoFCLznsrZCC2GRLACM5NC0d/N1vijpAP+8l26TTxw/Zla0WxmszgjUVRi9sn5nlJYaoLV2O75pmxqCyGueTuVZBR5sIjgzpJGXSDcb4Ih3ms/RNshlTtRGXCAYtcTQVB584ruFddwGxL2GbTKmbUSL/+cbmISMfbIAvAXgXykZA/yOpwL8NYzTAf45LtmoL1FgCHythyxXCalqS+fLzVFeQ9ZcwIoJrNgXc/XpxHrJf5565A39PApvrnBTj1EdBE0uHhTQlzVmz6vBG1/4QaCkgPPyDG6mPg04eMnwtMU71OJF99bLWrzYVgLzsnWHFE+5eTplkZ/8TSZFa+2gmbn3TIn9+7CVFeYJjy7UMmYs3LPhwjxm7kY12TVL7c0BeQ1cLADfkMUC8E8hDg4BnomoduyPfooP/hp+W7DlpQM8iZofJyh9qgyt4yahaeGEPn3xFsZbbVFAyefJ8NqA4wMBFWqDFIgDxy+bvfyAV5IR3HuNX6PcCNy0M1fXi5JuabzbYp6a8Atd99UarkkhZDw7M20au2iMiXJFwR00bR0OR8Ub3jRfeIjs5IVEjJ6zSYTTSgBUA8KGkEu9Dp6cJb85IWhpPZQ+cOOcqYppRXGs6vM9PsN5w0nUbZyHAsa6ruKnTKzecgzthruqS6HrkodXTeyHBeAtAP/SidofO0ugYGf/0mimSprPm4DqSvQVdf6mOk35Z0NOcu1QycvRkNhGPW84U9x6+0ci+e6ToZOMAadriRuNXHiYjfADARb2i0BC4aKr4N0Wc9SR/TZDlwsARJjdsMvJK8GOkGh1+zyB/ZO2DQd8xkSNmYLBvv3p5xn4QTTpxatC1F2p5gqFMOF2BpiI7e9NZ6vJwr5xDMknChrHgqDKn7ROqPHXVcg37mFwISNfqMlTPs2191mIKZ852f7SaAZaDnJWG6r6oTXmjeHFy1yMOVZ8tj7WA/nEScS+0W3Hk8d1TUze28obe7gYsl+UAWOL53WRBi4WgG/IYgF4M6QDPLVwdpb065yNiLmSUqcA3LyTDUevg1ULAxnKDpr7xosSNVwKDoWDk7SHaG/Myc7Y6NLymmGN6TLht+w7LSC/Vp1OJfjp4E5iP8nzrmNWVh3xNx5Y9pkRO7zntZ8sFAQSDpq5djUEsV/8BttF64nx/AwHrOuawpz8Yul/uIosocuC/dIBnnWyf/w/FzimJmC2xqfdoM9UwjxMxCghjil59iJASLBmv1gX3U42bvtwKSHV8DWoCyfYVz7DeUDZeRrA8+9cbPk8F3zeBcr7AOrKVcTx5AY9AYW8flNcNRq4WAC+IYsF4M0Qv8/UtxQUHgvnzTpbA0/XeU0eC4VM0+K3qo0shrLxAuCXaQJzgisfbGsrNTEY5hd+Mv6JS0FKRSM9dvo67FaGoP9kL3XylFfCEbyY/oD+6+DwWGTk1Mw7w6gT5r6xX71fafeMLKEAmGtLQxGP5pN3zQc4o9vYVXBdf1htjtYFWiwXriQr7ZxuDp7CZR3Mp6PXS3DmZPqg9QKp1wNhUVefmnqZri/G0Q+c5q3q0xZWjYzb/DTS3yGw9xR+Mxkcr3YzLrSa+C3exNRaeM3+U6PX3zetk8R8NKyzcZ+l6DZmFVZtOqr2XkQU6yzZecUq9zhvMuIdpgynVW00841XRRq4WAC+IYsF4I2IOVuoXREcJtv6wWfHCXWn4dFT15QGS99secU9pSGbIy4AWQIIvD+T/nq6TRhqSD/5V100i8Dcd1+UKEC61sq7RPtMWKN81gEHL6p84gR3Fj1nOsMFQyOvqDsoeSFG1PlEdTKSGjKHkxE4PBh0/Ox1rNwUroSfPGFOen7jZVglymqSeqmNMtc5byOi1n4w4jKu3rir+Eu3SZkJz7XfC++lzdTieeH03sMX1VF+25VBSrhZP4VbfUdATEtHsAgTrLeoSzW0ejWqqld4wUNHZTLezOx4RBZNnitoPZjavDYRjNtfF7FfHJ8mfR0xxc5fhWLeTM5CSWmlAgS2n2PDxYYbvEdPX1OHmDw2H0NvGctvZUypZev1KV5xzKVPjMBhOglecMJ6mfiN4/hEf4yI32K/uJhHnElQVhtdUezTy9pTqQ9p4GIB+IYsFoA3IjaG5jEFbJzVFnWYaPWWo9gok4chgwQNTiTGLJsjbnZS098i5LX9uMrix5zjLYcsUwDPus19t6GIk5PuGjL0Z7EcBk5ZhwWugdIW5iy5ViNBGQGRi5WpQBNg/ILPqsupefiGWjvBhYNmDDINTTpoUQZ6CagtEE18zbYIxW/yk2BfF+/15xgG6CNaMMdt7rIA5Y7hTUk6wFO2+G9unFIj5w1ZXMjN1em766QCGt4MtVEmCn38PFSkFqJnGEtt4VqIxgLw3IT3lu+xTvr4jfvEb20P5rdOqTY5eh5ULhtaIcZWUxWvpE7e40mrxUP6y/BOtlOTU9O+1CRulG+XceZPa5ER3TX0LP1qaOK3LQBvAfiXStTw2DCavXTPNBbzleFyxqZ+bcSJx3fp527cd6lK9kVhpQamzF8hc+81JNGFQEuBm7zcRGMKAQoQtWKGcdblkqDGx1uOCIqMHqG/mpt9yves+GL+mw1N5B/5RqvhJwFhuiHIW3PPGhP7TvdOI3lHz7j4nRo3+rKrJ5nuZmG/yKsvO9cNanRfUB5IHEvT+upDfF7JhsiSqevImPgM5Y1y10iIfefva/ve9z2kv5Q19lfqZf/NPfcEcTzl58/SJxLfe9Y+NTRp4GIB+IYsFoA3QxoYa7fn6Oa4fnlGfYjP09XD90nmvvEyie3nYLENbAtPO/L/9LNfv1339Xsx8SkYPsNXjQMB/kvywsw3Xiap9hvxXuO/eV6bkuK9EN1trIN1mfsGiWPMZ83VU5O0Nmj1vch4ijzI+1p/avuu9i2OHfv/tO9Vj7XWl2fhFalh+tUwxLZYAN4C8K+EFMBIexQZJlv9yfCeEOsxV//LJmpnJP6bYZCkXwY5K1dTbdkhWXjghq4DpmZgBAstAQ48B4vjYkqm320I4oSo4v8z855Uzfu6AJ5U9a2nUsMAoQbGT+uT1n7+fNr3qhbDF+BVQ/SrIUgDFwvAN2SxAPy/CdE/T0aPnb9JbVzWJsicKIwc4elKumoI8Mx1Qm1PB5FqQJR/vwHAYKF/DdLAxQLwDVksAP9vQhw8Pd6ZG8inY27XCfKMtuFmXKfR7iqfPC+koE+fg8WQTONkWKxTtxYsZKHnJQ1cLADfkMUC8P9GRGZTI+dmIaM5GCJZm7uGA8tQyUORV7F49X4MnbFeafTc+PxR3ucGZosBzmoyMTSzLiG2kIXqQxq4WAC+IYsF4P+NiHylq4XaN90r46w2q7A83hzEHDu1DeaN5CwVXeOzM0rdE+q0LhTL1x9Wk2n07I0qJbFmHZj/roUsVB/SwOVfAOBv3lW3rj0XwNeRAvx5igXg/82IrhQOJAeRB34Y181Y8JBjl1QuGwr4w0eP1MAyPzonjmnhM6SConK4+BxWsfK8BcniqrHQi5AGLm8/wPNGrucGeIsGb6GGIA4mT1dSo/+u+xJ0/nWlOtjk5BWqcs5Q2M9dSkJsfAouJ6SpbI2x11JxJvaWOgW7dR8PQ4Wi32QvFbfNuiwavIVehDRwedsB/hGu38rE9CU7lFXLNtQH4BnlxsNmuw6eN9TUMOW5AF7abQH4t5x0TZvCxQGg24b//2XQMvSZtAa8ro850Xl6c47TbswWmm6/ExNlEeBkYAZLav9Py5ViIQvVlxS4vO0A//CRSvvBMyfsC9vAb5j7Nonto3LEHE9UsugybcjCVCvMhGvR4P/NiT55ahw8ucrUwXrKWgqGHkfP063vNZ+LfzTT/q7dXUoNxXydbzOxT9S8dDL3jIUaljRwebsBngnxmLuIl69wTtFCJsiZ+zZJAbw8w2svmSRvS+AZs27R5y2bRYN/VoBne5hD62lpp03LawF4TlSexuSH6iTpHEk9b+7vZqjGN8z8nfQ0cFDfrcd7/L+5554g9uE53tMArfp5xrxzoLmaV9NCpW3weD+FV29jvdtWB/HbepuNqTb+PEHyHNuhkfm66kusg5OBCx7lkX0nP/R6jdtk7t0a7aqN2M4uNWWIZPbZZyS24XnaZI5vGj/NPfvkd16U/hUAnoBIzZcASYCju4Onqs19m/QkwJ/Go8d1Z019lsJcR7zchQoalbE6AV7GlW3hvx08Dyj+cx+uvuW1afBfyUc4YemGYDKu99VFBxrx47y4gt9WE1cJs4CbCBpXPV5ObPo828gJz3weqn75ScZo9Vc/qy5UYP0yyJwM7HCNdgmxHv79g9Zau1gHAYWM4TO6K4Xvs43M+c02GH+HxLayzUqY+C3De8yzouo38w6/xb+piyUINsJ3tlG1U4GNZq5xM1bltTG8w77qeXeq2qaE9Mm28f98R7cCPpJ/sz6Og+KPovlqIrAuU6qqW94lL5+o22Bh8Cf//t4vc5V1wboVT+RdlXelHjLFRUuNubzLfutpi5l/hn1Tidyk7TwXoMbfTB1srzamT7aX7+vja7xAsm2UO46Fzmf9Wb1vNfkq/5b/s63q7+Sn/F+9Jz/ZVtZtPD78pnquqo7quvg8Zcs8wGsyajy2rIdjore/oYhteNsBnu0oKmHwwSGVR0mbX9IOM98msX0EXj7HFN48e8JbtxqicD9g7fbjaCngWh9+sJ3EEMo8Ey2q7LJvMsBz8H7osRgtBy1TF2KQuo1diS6/asQLLphvnJsbXX9dpXJwN+nvgA4jVxh+v0o9oz3vrv7NZ3v85oHOo91Vcqdvuy5WAsVLqumX5jt63SRucNBPrZ/8ZLv0icF49A4j3LS2SZ18nrnCu0kdbAsTUpEn/EnfOOvlN9g2rV3u6GL4jt4/agFsFxc1foP53ruxXnlPe8fQdyH+vrvUxbQEek5wXTsjuJFv7I/Gt1WKmOmx0yh3FQ/PZzmIzCuvt6G6bdU84++Zc57v9h6/Rv1kelr2mW3g91mHlpO8WjNkojAKhtZ+beyM+8BvsV+9JqzW+Cb/Z/9bDV6mknXxfU4ekr5gmiMuVvw7wYX9bj/CVU14poxmNMS4+VtUDv1Wg13UWDDtL8eHY8L3dJ6RF9o4uRvaazRO5Ln8W8kPx2mUlrOe7/F9JhVrMchZjSd5zv4wVTB5pY+5zlOtLq3v5Kd6Tp5nZkn6cTnmrE8Bi8wlfXzIJ9Ue1aZqYt080MbkeOwf5w15wraRj+Qp283nKC8cO/6OMqOPVUMQZf1tB3iWR48eqbz+HAdeiv60dnCc2G/yf/XWYyivuC/9MVT2nIXvM004L7ZhveRHnQuNkG61cs4xkyr3E56lvFKAp+BRa2GWPV4y7OV3HEFHYtWtOaHHr+CgEP/N+G6/oHMq3G/iwq0qVIk3+PP5PaEX1LOMGuHz/PchIV5bxwswOo10xw/d7UX4V2PlxnCEHL0kQnJNPcsDQszeuD3oLEbN3qgiVahdsW1kNNvH8ELr5YHYuveMPH8FYVHxOCRt4ir+u4ALMwVSK2N2xVGzN6j0ATv3n0co+yDPHzh+WeuHvMe++e48odIDEzwoMF8Lg8dbb1X18b7Y0ONX1fMaXca+sBgwxTHzifeUSfu9tJH8p9bLiT16zkYlcMyvzv6ERV1TbVzmfUgBD58lkPDSCFfhH/ly4Nhl1Qf9O+RvSHicfOuSuvJw/7E4ddFGxJnr6g7SI1If+TbPOUClOFaTvINYMMIfahI8dbtStIk9h7Q2sG693/vlWyHhrDdW1R1y9LL6yRS4zt6h+E36TsBlfyi4pjJComBzPPh39uU3+R6jiKLjknEm9jaOnklQY8qx9NwWoTaZ5zjuVjIycpavkjESrR0uDiNnbVCRSNwsU/Iibdbby36zr0wHbOserECZ3/6w9QJ1ITjrDJXnyRfy57A8zxzvwYpvWh2Kp1Lf/og41Vf+nd85Js9HRd9U/2d+d26AU3aY4Gz4jPVKnjl2JE12DHIg48PUxTTJublOhUj1R/hPvhHUbdyD1DP8DtNO86fDmgPoJQsKn+EkNsfbZyU19v8CAM8ScPCCWoBp8ZFHdbWD/KPmzPnOucWMr8+iOZsrdK3wkvqlMq4/9FyiWRIio3VZEuQ9+Ua53Psc4ZqvFuCFqDXzlCUrZ1QIgSYto0DlRacZxZzovMhj+fowDPnDW1bcZWg+wElp3bz6bOf+aHXRMy9H4Ds89RmXkKouweDgUVOm+U9NjALCXOQ3kjLVsyReGsELQMJPxatJRtOWLhZdk6cG30U0MbuVwTgRfUNdxnBEnp23LEC1gS4TMp2MIgBQi3T1DUPinWx13ZzWjwpp330FAH8s2aHqIzBqLoBFsoi4YokA+DkBK7afbeJFEzzYdCbmFpbKRKUGzVWe3+EY6BolNWEuLH6ySKWk56t3OPAMp6QWq/v0yLPB07yVtsBwSk6SQuEvNYiktFwF8svFZHUUYSPv1stCRIC5mZytNqV4AcUGWYSoEVLQOSHIH7pGaEVwoWF+ffo29Ysx2PdLCWkKSLnYcmL4ywJz8codNWa8sShO2kKhGyFjQzmgkHNhNZYv8omAxlBP8iI6LknlxPcPPqty0vPEL2/JYrvZjytSZ8rdPATK4jhu/malodFlo7us2N4BMk5LpE3XEjPUWQElO4oflYg8d0OdO6D8sL/kN90ktIqYz53aG9vPm8FoXtuv2a8UkDjpK2VWyZb0n23YvEcWMln0eTEIQTu/UGRC/hYqC0AnWeQJ8J91XKg0+BlLdylA59gwLTRlgXykGU455y1SjftRI9cAR7dqeCMZU03TZOdcIehyIecVlgRSXVnR+fkiRF78qwA8F8KxIh9cvHlRztddzKeHJpHPX8hcotwvdNuHZ93cNFc4r3gL2nxRGlg/+cqf5r5PouySZ1yQxovMU06ftbwWH7zeOa5ivE2eO9zGhXeYEvwJaPw2Bfx96WTrIctFYw5VgGJcIs5eVwBH3ycXEAq3WhnFtJ1ks000qQTDk9WFF2tQU1YrunSW7VETSd4lOFAbppadV1iqLm6gBq40I8PE4bPkDYlCw9h04+vr7j94pLRLLgJ8T19A2HeCGs14aveV92r2hXHtM2Xis03aRK3Ju4/baRr6ig3aosK7Phk6SUDXfLDac/Sfc8zIR2q9xoU5bTz9ItB7wmq1m9974hr0FyGYKLxaL30mmLIEiLVEa4u8Yd2sjyDDMeFCOFkWXNOLwM+Jlr3QLQhthy1Xrokx8zYr/hk/VyRARquHIM+6TfvJ+ql5cxGJuZqiwI9XFFI753c5Bry+kZosFwBeWM1yVvrJk8D0j3MMKcysT3cHcTEmL3grlF7oW+UF6PS1fmkYTxIBnrJB4ORBGVoMNjLR6Xqh/5Qha+elr8aFisocxwDQVdNDLMhpdv6IFA2eC9vZmNvKlfIBL1cXXjIFMS2ZRaJIpGcVGGqoLuQ93TYftdX2cb4xAiMqJVz8eMPWdVFeuIBT/mnhNbQfnrzgGP0rADwvdadlTH5qioX575MUhkhbP5Txmmq3XZSIu8/sHjEtlDXe7sWx4hiam9+mREwj/6n1X7tVd1pxc+W1ADw1bLUBJ9rMjKU71eEdvXAT4risVB1GrVDfZQcZcfOePMufkxZuUxdW64Wr6oaAkwpM/tZklgJ4AjsFnaGF9F/yCL+5kpZZoG5WohviE1nRFRDIwFLLomk2RQaWZvlYASkymZqhrm0S6Mibd5rNVr5/nnajNsdCoE9Oy1Nx6gRkklpwhLkcUA4u+zVINGxe30cAYeGEoEZmtypYCSC1WH1hIPF9DhTb4bTuIMJOXoXv7iglKNzI1K8f5HPkHXncZpgr9kjbjHfeeaqPt1rxmf/5eYZ69+9NZ6t/0/Kxl0mgX+hNTVPX4HUZYL1sP/dIaIXo7WfZLxr14One+NPP0/FXGY+/NJ6lQJm3IZHfeqG2yrYTlFg/ec+6yR+ewKU/mbcgUXsmQK7Zekzx7a9SHxcvWl0cb7aJV+5xAaHFpS/0NQF+Id5rMVctVrQusnKLDa2AsoJ4OIy8Vpac8JLECdB30lp1GxMXb2XtSb//9NN0/H/fTsUQGbtQ6ateOPk5eXk1IzeU/+vH6dK+2codR1fWsdMJys+vgxz7S563GOiEdWIVGLeJhX353Xabkh3Kgm7GU/64+BAwaU3QSnD0CpWFg5qephTwuYYitvVfBeBphVGmaU1zw/1p13USCIkF9CTQjWmqWD5roVIZcjROydU/ms9RY2nuuzpRRhjJQ8ucQM32P2t5PQAvxDpJBBreR6oXgiNN4a5jVypQ4TMEU05uNm6CAHzK3WptkAC/ScxibjiRaQrg5Tk1SUXgmYCLvsycvBJl/pruhtN9s9hjv9KICAJ8jwDB/9P/zwH5de5GxRS2x3hQ+DvmZ+fE5V2qNMVZ2IckPe5W6qPGxmf1Scp/v9tinho4AghTARsXLhZ0DzAlMIFHByq9bc36OSkf946Qc6ptjGYhn/R2cawICnyfG3X75FljgKcmYL0iUD1LnrFdbYcuV/sKXFAYwkW/sY9YGB0EaNlvY4Dnt/h/buxFy+JsDPDcixglmve7smhwEnHxoP9+8iI/BZTGheNMNxIBWG8/v08AGzZ9PaJFLh7K+HIxoCbNxfqvjWcq4aR/nROA/+YYM6f+nkMxmOMUoAG81GfMNz5HV83OA9FqwdBLiiwM3OthxBKvDGS92vNWGDTVW7mZ1myNUBYSFx7WzYWLbjLuF+mFAH9ZtLzB8o5aUORZ9oPfpEnOq/goJ2wz55S+0PNbumVh7OOl5cgTlM36azmFWCf7wuPz7BfHPfpyklhHx5SssL0cEz7TkMS2vi6AHzFzA46evt5gAM9CdyGvzWQ7iA91tYVjxDGkkkOXY1Fp3ZfBP60UFpVjkXuwUjS0vRjz3yWxXWwf20lFkFba85TXBvCc/AQvgiDNb70QHE/KxGfUAldYPsfnGX7GCT5RAN7YpUOA37L3jDKnqTkZAzzf4QTj4NDUpyZlyigCPv/GeulfplXB97jK83cEOro52AbywtikIqMoANwACZEVnr5YFqXBSxt5LycXBE48Pqu/x/Zxc5dt5Td4Y5NxYRoCbu4poDEsbHyPA073zK9zNyl/HDdR2U6l6RsGT9UvpOeLZ0RGkPTBGOATBOC5Scf+/qP5XAVesxx2Kb8u3yffF4kVsUAWgfYjV6i6CTC6DCjAkd9x85tuCmOAPxR1Fb8Kv943aJMEMWrT3PTl/oJxoQ97su32qud0q4j/5sbqrZRsNb6k+MR02Eh7uPhQ02aoLOWSoEuNnnsdy4Uf0+13yd9lcZPf6e3V62cEBUGTF6johZq/q1h4fJZ9JMBzrFgnrUIuVlxsGE2k95uLIseAm6N60QF+6B8+aszZNm1RtFOyyfZxA59/060ytktfiFZtOlrlGtNLwu0MTBYrkhFflBdGYVExoAZIi4gXgdNCoJbHb7Hu2jbsnpfY39cB8FQMhgrfD5+Ib1CApzJFXjcf4Kxk33hemiN9btGlysAJ7v09T+EeG4Mduo0RXJOxYr3mvqcT2/X3ZrPBC4EYVPEsfDcurxfgZcIQyEwBnpqdAnhhhCnA00XDjQq9KIAPNA/wNKnpC6YGz4gVArXH5qM1NDi9BIfHqWgLAgFXV2rwjJqhaUYfO9vwmdT5JMDPU+FzpgDPRYgAT2F9AuCFOGlIvwxcphYebi7rhTvtmv/bQfpAbVV7l/nfuRewSvqwXyyL3238DH2uaZbrAE/LgSBlDuCtV+wT/tLFNEf5snkxOTexuS9CtwD98hQMai8a6FV/owrgJ5gBeNFqRwuf/ykAzH5TkPj8iFm+ykduXLiwsQ+sWwG88Jb18t8EULqv7hv85Qxzu3ojTUVGUQMj4LDffJb1cwOTPvshArBKSxZ6XoDnOwRTLp4qFbPIgvqdENv3DgFevlUbwHO8CLq6rDCVc3PhI3mr6hGZ4E/KKTf7PpDJN2jqOmWVGRfKU4DIRjfR/DnO7Cc3BulHPhF9E45eB1Wf6Ips6PmpE/v7qgGe/CHADxSeUMF6UdeIcbkni0XctTSlQLAttMiM56Y5IvZwLOkKpGLFuWpsbdVV+BwXqCMn4zF+wVYli/ymue/oRBmkVc75Qy8C9w6M942epfzLAzy1JoYyRp69gbbDl6PTaDcVMXLvQU2tgEJEZtDc/buAHvcJCD4vA+BJFGK12SO8tVsV8sQGCl1HI2auV89xo0cHM+6m05/LPtFPzbtddcDQ6WkAz0gSug2oiVIznGC9Dacu3lJaPHn3ibSL7VVjLv83rptUF8CHnriCUXM3KhcFFwqCIeVnuv0OFe1kXCKjbyjh+1y+x2dYtz65aRVxEWOEkV441nSzcWGg24TuHcoEnyePad1w89W0vc8C8MaLJUGe7yrQ57gLL9jv+gK8Xg/f53vGsmP8N7qG2G8t10iByhzKwv5m5hRivli57B/lcuTsjUohoMJC9xt//zRt8EWI7X4xgK98ToCfpyxK9rUhAZ6FfN0r9VJ+OJaUZ3NjoxPHk5vibBddY7sOXEBeQZmhtrpLbkGJ2sQfIcqjrniYYoEx8RvkE+cmrd6g8BdLVfxvAfAu3odwQuqk9suJ3meip1pRGZ1hXFLT8+HtH6lu2ef7DLl7WQBP4mDTxUCNNCD0onpXLwx/dBHg4UGevzeZpfpPbdpp7UEB40QVgUJtjhuupvU+DeAZpUSrhlolNWX66JNSc1USM/ZbRXnI+/y3OcGvC+DVJquA3J9+moG/NZ2txoBAREEzBlZuuLptPKI2YBWwSn16/R/IgsbNV+7PMF2yaaE2Q2CmNkWNnlq7ZtZrbhv231helSzUF+CN2sHJxnGrGrtnAvjqcWHbauOl+rs8/16LeWpR416DcWQVx41t7iMWFdvHY+6MuacWz/o07bJ2wHhRYn9fBOALi8vUPgatIMoN+cDxMfctnTSAn6vmKfdMuNHe0IX9IK9pqXKsKDvK6pPvm8M6fUypRDCSivt2HBdapUxkxsAPegYYTsn/M3Bkd+h5FenEb/C92qwF4hV/cnx1S42LG603htm+SPm3AfioC4lqAjGygQykO0Lb2DJUZCi37mhhhwRWRk7QRGQUDdvQkABPfvJ3BB5G8VCLLyiuvsmJZiAPyrANjBbhQsC7WzfvPSWCc0GF2HGQWL9xvapuoboAnrdC0QXkuDZUATK1DPomGZ7JyVzbgqRTXQAfdT5RO9gl1gVDBWkK66Gw5aKJMcabvmaGYw6Y4qVky9TFxO9zsjECiJunjHQxHSe9pGYUiDYWo0I26Vel5kwZoDtFr++ZAN6kLTWIz0hdzwrw9SHu5fBcAMN6k2WxNS607sjD36w2q6gubn5TWeE3OF46QLwMehGA59zMELmycQtSPHmaXOmkAzxP+DLSqqE1eL1Q5rcFnVVySB4SY/RFyNxiTIuOckTigkUQptXLvZ91ohiyrfzJ/zNzZf/Ja/FTb3v1fG0LMX/H7/G7xCtiHC0L//3RNeT0ecu/h4vGAPBceclsTia6R9TBFxPXCIH1lDw7w36nmsSMaKGWS+2toTV4Di43NOkmoRbKTUhdmAnI9MUzTpoTjH3hoaW9YTGYL/USxDlI5sblaQDPcE764a/eSK/SEKhRcyOQbSWZE0ad6gJ4WgJ0H+wWweJJ04TbmepwFQvBgYsWD3kx/p7jSzI3mdhngjUPfK3ZFqFAniBqrlAG6HaiVqvF7Wvt03nzNgC83kZueNN9YBwSx4WR8sTQSZ4Spi9XTVhZBOsap4Yg9vd5AZ4BDJdv3FV7WVpQQf3ay8WZYdFUYnjOgW65l1VoDfKwHM+AfNdjsZpn5pQOYyLGUDY/EEuTz9LNxsWZUWjU1JUPvS3zO1X/3Vw9JI47rXB+l5vpVDy5j3X/4fP53E3LWwfwKormBQCe32R/VPiTCJDD2oPINolB5qYIj6bzIBRPR3KDkG0gLxoS4HViuxntw91901hXnsTt/psHfhm8TMDpgjpBydBFTm4C2fMAvNKgRQscb7VVHca6mZyFzNwi5aKpAfDSV7bbFIDrAniGofEU5p7DF5/QQOga4ulZtpnjye8o37ZR3TpxASYo8FsEZt6nuefQBQX05grlgNFQXDx4KIjv6dbN2wDwJPKEbZhk44fTF2tGHOkyNd9lr9r4VVE1LxncSezv8wI8lRXmf+E5EVqg5uqvjTh+DF3l5jN5YSxjDV24T8D9Ic4J3nPMzXBaEDzjQiDkPFPzQNqjzzf+n2NFVyb/Th5xvimS//P3/Ls+7zknOV4Ee2IQx48BDgwH5feGyXc37D6lNlRNXccvUl47wHMDyRTg6wqTNKfBbyXAC8g+DeAZt81vEjw4IIyWofBtFcFlnLxxYUz7DdE+eUqWAE8+kBoa4DnwBB/+ZH4catXG5bRo9c5eocofTbcKfedsOzd96qqzzk1WFUUTqOrhprK3/wm1qcsYcpqhqr1KGCnENftMqgvgw07GS78DlcuEp3KNS05eMfxkwWrW31FpNxxTU7nitzh+PI5PINMnw7eiGQ2Wyc67Z3m8n3fWmobPUWPkBjT7pC2AWlz42wDw7LO+6d60v5Nyw/DErz5s9GdfF3lkzDxj4TkG5uppaGJ/ycfaAJ7zj+NAMOdBHmrbdH3wMB3llftLrIf8J7jVlyiD1Gx/7LVEuUHoKmXYLP3cdPPRquE3OeZsQ0MUbnDz4Jiz10GMnrMB7Ye7Kq2c48KMq+yDruFrOMLwVA3gOd84Jmreqbln+Cm/4/N8j+9ri7gNfuppr9yYjPxykvnNvFcZ2TXPwzREeSMBnho8szdywhkDPM2g3215PL76RGQ1wGsnCM0BvO6DZywrv8l62R8ynP+n9s+kXaYrJ9tCE5OHWjjxyIuXocGzXgpG035OKmqAB3v0kp5ViHOxSSqh2lq/CLWzzgNZfMdcXaT6ADzdPP8ULYWRGUNUvPFV5ZZicjO+R/7x35zYDPMjT/W+1wXw/FYfntSTNnJDmFq7cSGg8uRoMx4cEt4Z84X1M00yzVzm72HOmeYDnRTA8DwAx5f8ZBoJRhExRw3dTUZdU3znaVamOuD4sl6mIOC/33QNnhoeJz8nH2WHbjQ9HI/x1ycv3ETnMStV2g6Oj7k6GprY31oBXppWXFqJy6L9no29LYpZokpUt12sTIbhMmUDw/3YH84B8rfeJN+lnHHc+X36s50EeHeLFUuli1Ff/GZSak6NSKuGKPmFpcpNy/lGlwnxoXl/ZzUftMWnGtw10qxeY9IBXntOU5i4WDCwgedHeBMbAzr4nRfdSK2rvF6Al87z6jlOVL0wVQFjfLuO1QFe01TYIDKIjDE+1k2AZ4InnhLkwYX6Ajzr5IRS7ZDBGGu1GeGnrxlq1QpjsBmLzVWWPCAvXgbAM7sg66cgW7sGIja+WvPlBKemclu0F7pQuLnztNjdpwE8gYMJlNh/molMosbzAXTZcCJzXNgeprqdZkiWRsBkP1h/XQBP/+HQmevxlyYzVR4e5vMhOOlf57M372SricO6CLxsh7aAaABHgOfm7DZZuPtM8lQ8JG8I/krLpUyIPHIziot7nkwQvX/c1OMhLWr/+mlBTsq3A+A1a4OyzrHmQTC9X9SOac3RXUfNlu0wV0dDE79TG8BTNrn3wRBYWsdaymQtzTSVMVrg7A/HlXXRt/4spN5RMqHVQ1cGU0J3G+uhZJLRKZ5+x1VOpoYs5Dl98yUyn+kt4EJC9+hy78PqXle6U3rInG8jGj4VFfKGGvmPgk/8yYSH3EvhoTwuDsNlPjAli6vvYWXVnL98B3czC9XCxD0/47nZ0OW1AbxuutCnTkCvKtLXMzG3FcBzIlNz5/PvCnhTaGY57FYhhMaFaVcJUvrJNB3g+a7uomGYpO6i0QFeTWo+J+2g+4e5Y4zdCmQ+XSY8zci28tknAd78SVYd4CmgTwN4AjJBle2lQFBLMS4Mmzt+7jr6CtipzSpph7l6dNIBnt+t7SQrAZ652WkNUEB5yIsThzzmosWTm4wEYCoF1sELS3S+6QDPMLYnAD6KqQo2yFjMVn5M/Ug/eaIXgiFNbn6TJ1LZTvo3yVvNJWSrNnwZakafMzdOqb3z7zp4so1c6HjakfskuvWVlVuEBa57VZ8o1GzvswA8nzXmZQ0SnrDfBHgu+mHSL72wT8yUOeQPb62NXITM1VEH6QDPts52DMCla8YAf1/NCyoTbxLAk390Jy1dvV9lYGWmTf5kcABDBGmtUcEicR7Wn6rfo+bOvZVFK4OweFWI9g1ZxLkBy9TKDE18mYUuKMoM3YLc5yFeMfpsx/7zKs00XZzMlcRIJ/5ctz1S5VFi/DtPr/J5pungQsR69MN7r6K8FoBnXcqvJROUoMUJalyu3rirJpA+mahlM8kUw+CYitXYlUJf3GIZbB6kYCeo7bETfJed4Oalk1eoikDoLgDGb9JsMm4P3+HBghZSP9Ma6EmxaE1w04O+ce6as05TgCdAcr8gODy2Ri4apirgBjK1GIJ3XQBPUqafoV0UYAqCPrnTMvOVkDfp56BA0/RdUyJ/CcIEeYbTBZrmohGti6GMbBPvciWo0Pzku1xkSMNm+GBb0BmERl5WgsAx0PtAeeCk52JEEDbO73Mw8ooCeLaT4Z08eETXD01p48LNZF5gQG2P9an61SJCPixUmhJDVikb1IA0OdBC2NgGtuW/f5yuxpp9YapmFk5C3jXABZtjzWf5DhdGHvvesd8kF006c9EcMlgJ1fXXRlyMGN+vFq5Ik1w0N9LUpiAXCv3bz0K6BUOlwZwGT7cEFR89n7mxLL4sosxXA/ypGosjC61LukwZ+03iON+5m6vOlKQJ+DcUcSHhWDGElInwOM78HvfOTPdiXlWhYkMsokuVEW/kDX8y3Jm/N1Z8Xld5PQAvxInECcuDLpxgzM2tYxAHjf4vmmOcTNyBJ5COmbtJpW3VAYVmFAGYJz61zQ5N6NkJTgBqfUzWRICneUufIDvK/pi2iROLmhOf37j7VJUbyJyLhu3nT77DjIDcqOVFHXo4l67Bq2yS8h6Jpv/TeMgJzjZrWfSuKQuChRkbedCFdZhruynxO3roFdPgMl+8McBzs85mxT5lzfy50UylDdNCekf6woFneod9R2JUdA21fx7Z52lavX5tEVig+HnWBOC5WTRytq/Ko0F+sj4mVfPyj1ST1LiwHdTIvhFest8cM/KUckEXTqrwsFCAm9ENzNPOPQO2VT0nY/E/jWYoIKeWz0nGA6C8QIVnGLhY6hYHf/I4Py0BXpJRE+DzDNkkF+Jj6Zf+jjmi24CnfP/SeKaSCYaB6oX8vXYrUyUoY/vZb3N11EWceHyX+xe0VJk3SffBE+Dp4+44yt2QyK3u05cNRTUB/vQTAM/CJrKdr4cMjXiNhU1gO6pI+/UbUV6bi4aThR9i3fSl0RWQm18dIkiADQi9qMw9ajOMAT95IVFpDHphelYe7+bEJegQIHWhp0bG//NwEP3C1PDoMqEVwD7R723aJraFoEfgIijyMghq4swxzr/xPb1+/mT7+Q0mfFLaltFuPkFvrd9x5RbgYqZMf+Gh7lusjbiQMf6XoMM6GElBPzb90uoQj2q7+XdJevuojTLyhLcvGadjZuHGJG8AohbKxYn7F3Rf8f+M0uHFCPQPZuYUq0ydBFFqwPo3CC7M4MjQRR6aMi4Ebd5exG+r8TVYTMx3w1uoKg2Lll4YGcG0CfTXq30DIfKaUTh63QQVCupIaR8XYPKAIE93EvcIzsYmKQ2atycxyorfJSjp7eUmK39Su2bmUj6rF/KYpnS7Ya5qLFk3F3D9XWOqGnP5N89QXDfZQOYCz4udOX60SMzVURfpY8dTvAxf1a0SvTC8lakquB+itfPlHXDSyRjgecmMqXvUUt7sQgxl5NUrB3idqGnyJzPkMe6ccdQE9wcPH6KgsExFYfBCB7pNSsvuKdOHgk4Qol+OJj4nO8161sPGs62MvqA/lH4x+tVpOlF7IyBw45HhV/T7ciHge3yH73JicnOOAMN0BlfE7GY0B4FHB3jtG3bKwuABJd4sdCNJ869xUnJXnJon/czcFKZriIDIRUX/Xm1EzZj10y1BM5QRRayD2jCB1dw7OultY5ZIuiOGztDaxoWQgE23CNvGmHdqh1w46OPkyTtusnIC09Lhc5WVDxATnypAdkAB2gfU4KXt33XX7mRl4jbmMectRtyMYr0kXkDCDSle9NF++Arl5+ehD1oA9KvzfAH9uOQTx5ILGPnEG5p4gxX7yDFgLDitJ1oR3DM4K1YM/Zr07TKSiod9uI8QGBarfJtMW8ANe+1IvGFzT9qtxkk0em7IcvGie0rng/op7WBoJS/qYHQDlQXG5hPsq/gqxM17bvJzr4cuO55HSLyTU1UX918YHkge0vrgHcK8ak/bW6h7zLW/28ni5az2NX5f5KcsJ84DvX4CK+WeqRloyTIXP60avscxN1dvQ1AVwIsVxzMZFoB/uwo3dfu9ToAn6VobtTPme+dxX4ZC0cdGYM4vKleTh1oyXSEMv+snCwI3AumG0AGXddHEp8bJaAP71SFqIjKnOzP1MQSRqTeZm5y73JxYxptq7CN/x87TLUFtiULNXXM1yQ2gwQWFWhTB3d5jv4rkYD74fQI2PPlK0CFoUTPkAjPDfpcCO9PvmSP+nSGZ1AJ5XRw1RbXhLL9j3+oaB9bPNjLaYMBUL+XLZ5ZIas5622iZaHelxkmb41Rb+XveS8rNUC5q9KNTw2fbeQaA9TIFL3/SpTZCNGm6vRgOx3er+q3qiVPa9grfIypfDmPeaRWRp8xjP2WRv7pXl/zhmPAn8/1woeF+jFpIhb+MQ2YeEmrEjEBwWHtAPUcwZgpZnj6kT5rAzvwsdI3wohf6p7nQkQjKtHwGT/NRC8NGeY55d9h/vb38Pu/q5Eb9HGftNiZ9nPhT46s25oyYoBXBBZEXhxjXo3ggssmbqgjCvCdWuwDbYL2ZjJVO+rcI1Nzk5r4Nb8DaITxU43NY4y2/RX7xbltuCnNBpCuNcs85aq7uhiDWzb5z3Kks0CXKg4G0rhraz26hhiGOCzee6epmWmlepP+BKFjK6nsdAE8h58cpSPwOgZ4mPUPO6GKYKNrcr3M2KVODGhQPwXBS8HlT7YV1Efi56cqDBHRBcKOxwwg38PZ7/p+RNdTGtOdrtoXEOgkOBHm2g4uBxhijbwiYMKaVp0pZL7/FtlWTmzpww+8xXIp1mX6nNiIv2AfWSx86k58ZA87TiBo8wYgpdFX/zbRN5w0jZEi0arqMlp8GfvHvvCT6J0O79W8r62DwMvUM0wiw/8Z18/+sg0Trinzmu+QpfxIoeICEqQr05/l9LoBN+zpq4yljwjTK3MRt1JtaquGuXHmOudmpvTNnC/POU+Pld3jbl+nhL44tx42b5x3lW+yfaXurx0lrA90j+vvGxLazL2xXZ8O4PtF3qUvvexuZPFQC1LsmddVG/Dbr6SLvs01sW4365Xud1BitRFvhIc8ncGGor1w8D+ljR3knUDDxHq3bSQv9VASchd5MogXMuxY4P6hoKOVUxpL4yvF8pQCv103ApkbOj9N1Q7eLPmHZQJrv/Dv9wWzT0zYu2SFqhPTP8z2a/qyHf6vrPf6NAM/2sP9qggpQPPGsPEeXCzd49W8Yk/H3npXYb36b7xuDVv3ITvlnq/ovvDJtG4k8ZBtJVc+IBqwT/84oINZp7JfWtQD6zGvrt2mb+Q7HS40h6zd6XvuOEZ8MvGbf6bJiVIk+5vy9ytMuQkvtnO4DyoueZIzfqarHQNqmuDaext9VVNVP+T4XBOGb6fumpPf9ibqEtL6LXMpztfny6yRpgybrMh4mfCKxv5Q5rZ9Pb2tDEr9JoOci/23Xxer/FnoziePDcVL/Nx5DIcrvKwV4U+Jk5HcJ8DTvKey66V5fTZaTgIBijuoLmNQK2Y7a/OZ1fcOY+D1zlsLTiN/m+/XprynxHQKeaVuehdg/c99mf8w9b0pav2u+r7lOanne6Fn+m/3X+KadoiXwEeAI5pRJgizdVk/jT33bS37VWZe04VnG3GwddZDGm2fgrZk6XiZpqUM0hYH8t9CbTUpJFsw0luk3AuAtZKG6iMD2qsHNQhop3stCpI+Bhd5Q4hgJmY4fcbwmwIuGRIA3fdBCFrKQhSz09hG9IB+0modWQ5bhP3hbOw+haK4SC1nIQhay0NtKxHHum/2t6Sw07e+A/+AGF4/xm3vYQhaykIUs9HaRAvgmAvD9HPEfs50CMNNhF+a57FExuhaykIUsZKG3lATHed6Dhyft1xzAf/Bo97m4JJy/kmwhC1nIQhZ624mYfikJlxLS8B+GE6+WYimWYimW8i9WLABvKZZiKZbyL1reaIBvyLScb0LaUUuxFONiEUlLedmlwQD+wf0HyM8vRnJKNq5cT1P+n4TkXGTll6HCcMPJvcp76mLm9OwilFTcx0Nj1JV/l5Yy6VghsgvKUGnmVpRHDx+qb6RnFSAztwSlFQ/w0ORC3sePHqGiohKZTMYj9ZRVVqevffDgAQoLSpCWUYC8onLcf2h+ivE7JcWlSE3LQfyNNHWTT0JyNjLzSlFxT2sXnykoLFVJf/JLKmqkqGWplL7mSVt5xVxp+f0nJ7P0t6S4TLWT/MgvfrIOlvvC15y8ItXfwtLKGvUwX3ZZaRmyc4uQnV+Ke9If9rFI2s6kREwYxYuGee8rL/zV+FaMgtJ7T/CN41ck/eE7GfJMcZn5m9/LyyuRK2Oo6jZQus5rGQ+Tap9StDG/m5GHhMR0lRXz6q0M3MksRJF837SNj4Q/ZaUVyJD+8JuKZ0+yTHh2H0VFpaq/7Ht+UYUaa3NNe0h+FZaocbwrz7If9wyV8nscoyyRSZ2X/Jkl/edYPDBqH8eipKR6LO4/MZaPUVFeYfh7GcqN5LJCeJqVlY8bt4QHCdU8KBa5MeVBXUXJvtSlzaFSlN17UNXnB8KTXJGjxKQMXL6eCt5FfPtunoGH2lOazJZUz8+6vi395XzOFvnW+VYpfDN9g/OkjPOa/JU+5UjfObfl9ScK218qPNT5TZknn0yb8Vh+USn9NP72k/zWLuwoNswxXVa18SsR+eL4GR40LdK4RyIX+YIVfD4jp1jhiGk7Hhn4zTFVdRvkg/zLNWBPHRx8JaXBAP72zRRs2nYYY2b5oMPw5Wg5ZDl6TVkPl62ncTmJF0E8xu3rSViz/hAWrjyAqIQMFFcJubDhXjmOR17EUo99WLXrLG6m1bw8QiQFeRlZ2LglDHMdd2PJunAcj0uXgarOH89SWVqKuJhrWOIWiDVST/TN6puGUu7chZ/fEcy0D4DfoXikFz25iLBkZ2QjJCQKk+f5ovNIV7Qa5oKeU73hJH25dDtPPVOWX4CduyNh67oPfsfikSkDalzi4xPhs+UQbFYfwsm4NDwwGemKsnIcOnAG9iv2Yt6KYGw9fBVZJnXg8UOkSZs9fA7AwfcoQs7chnGLK2UhOxV5Fi5eIXDbfgbJhQ+RnZmDg6Fn1C76XKcAdVkG74FltkWrZQFw8D6CPaeSBaBq8i3lVhoCdkcIb3dhkVcYDkcnyQJs+GNVeYDz567Cc/1BzJG650p984Rs3ILg4XcKRy/cQUaBdilKfUplaQmOhJ/DQid/9BizUiU/a/frKvzuGIjg07eRLwuRcSnKLcDJ4zFYxIiv5UHYGnoZWaWPnmhnSnIaAoMipb+7pY2B2BoSizs5FWYBK/NuhvT7GBa57sGcZfuwxl/4mKldil5cUIzw0NNw9AjCLJE5LdIsAI4+YdgZfgW3srVFlYUAGh4eDVevg1i57STuCh9q4MfDe4i9cBUr1gbLWJ1D7C1NjvDwPs6dvQIn9wD0GbcSrYe6oP0YD0xy3IvQi6nIMxmnukq5gOPlCwlYunIfVu44ifPJeYY+P0RS4h14+gRjxFRPtBm6DG1HuWHYAj9sPSLzoESTquvXbmP9tqOYv+IAIq+mo7DS/PxguSeKxfW4m3BdE4w5znuwaudpXM8U8DZ5pSA3D1EnzmPJykDMcgiAx7ZTiEsqEKXvSXQtKSrC0bAzcF69DzPlWQevo4i5kSMLVc1ny6SfsdFX4bJaxsVJvr3jNFLyymvyW0p2Vi4OHDiFxe771OU8eqSJk+8R7D1+DcnZXBielIlHouxkp6YLnoVjjoPMB89DOBOfheKKmp0rk8Uj9lw8VnjuV8+xbs45G7cQrNl+GscupCCnxLyi9KrKiwP8g3s4EXEeM6190aiHHRr3dUGviV7oP2k1WvZbjFZjvOBz4IoM/CMcCj6G7iOW4+fBHth3NhkFBnDmyl14Nw3LXLejce8l+NXxAM7fqHnRbn52LgJ3h6Fdf3u822Iemg7zwLId55GaU32JCEvmnVRs8Q3ET10WYrR9EI7EaxMp404KXNx3olFPO3Sc5IutYddQbDJ32I6ctLtY7xuMbsOc8XM/Z/SY4Im+E1fjl/5L0XOmHwIiE9WzSfHX8dv0tfi2uyMWbz6FO1n6LUIiMI8rsHdvGHqPdkaLMT7YcyKxhvA9rKzA5YtXMGKCG75oOx9f9XTE6CVBuJZimPSGcr+kCFFHTqJz/0Xo/PsGeO6vvkKOpTAvD0uXbEDL/k4YsSgQScWPkHDlGhxdtqkkRI37OqK1LLTMdtlyoBM+a2uFliNXw3mPAGOhflO9tPd+GbZvP4QeQ5fis3bW+H6AGxw3n0TJfSOt9/Ej3C/Ow7q1AegyyB5fd7NHk35OaDHACU37OeCn3g7oOG4trD3DESuLYKWZCWxcUpNTsc4nCN2GL8OPPezRepi7yi7afqgTmg10xki7fYg34ceZqIuYNmsNPm+3AB93WoIx9sE4dVu0PJPV86CA8qjJq/BVJxt83MYag+Zux+4zqaJR1ZygBVlZ8PcPRYcBS/B9Nxt81cMRA+bswvnb2lgm3bqDP+Z4oVlve3zXfYnKOtq071KVO73daA9MXXEQZ65lqmfvlRZhkcNWtBrkjMHzdwj4l1XzTkp5bg42+uxDq16LMHxxME4miHwLuJ+MOIfx0z3RpJfwYLgbBk31QrshzvhliCuW7ozBzcya8l1XybibjU0b9+OnbnYYauWPiGvahSVXYq9hseNWNJW51WrYCvSasAZ9J6zC9z2XYqbHUVxI44L2GPsDj2DQuBX4duBq7IxKRl5FTb4al9vXb8Nl2Vb83NUW77VeiO5TN2Pv2bvIL6t54UuMKFvWdr4yxtoF7B3GesEtkPJXUxG4X1qMqIiz8n3Bhx62+LyzYMcobwSKMpJXXnPcEq/fwgIbH+mnLT5oZyPK1wacuJaN0hqffoQ4mWN/zPHETz2X4Jvu9ioDK5PY8e6Gzr95Yr7I6rXUwhrjxFIoytv2LSHoKrL4nmBNo4HLsTb4MpJztYVfL+miHKzz3I1W/R3xeSctSy5lhPLRqI8TOo31hP3GqBpK5qsuLwTwD+9V4ubV65gy0xNfCnh83nUpZjjvx9agC9gVfBrObjswwmobNh++IuZmJdZ6+OGbjtb4ot8qHInLQuk9jbU0dW7GxGH85BX4a+OZGLv8iAIJ43Lu9CX8/sdKEZL5+H+/noxPOizChOWHkZBa88ahq7HxWGDlif/6birGLQvFuTslKMzNh5dnAFr3WYwPOy7BvHXHcD4x54kV//79ewg/GIn+AszvNJuHEba74b33LHYGnYaLRyCs14ThaEyqNPgBToVHob0Aw19b2sLZ/yLSqgb/MR4W52Kl21Z82Womfv7VF0Fn7xj+ppW02ylYu2YXvuq4AH/6YSr+2dIKv4xdj8j4DMMTWsnPyMC2DYF474ff0Wa8LzYfu2X4i1YyUlIwZJQT3m9pjSEL9yKn8jFSku5g+64jmGLnj7mivc9z3ImRU9fgk5az8b++nYZfRnvBLzKpanGli+Jm3FVMFiB7r+l0vN98Fv7awhoTnUKQXEQ3mnoMj+W5nORkTJu+Ch83n4mfBrth0qIdWOi6F1NsNqPdwKV4t9lsfNvHBQ7bziEp68mr3/SSl5mFbdsOCtjZ4sM2Nujxuy/ctpzAvsMx8PTdjz/stmCCIxe8XO0Fuj/ysrFq9W781HE+/tFyHv7Xj7PRVniy+/RdI81KGvuwHGvXBaJJt4V4p5U1/vPH6fh56Cos9o9Rrqnq8gjHwk5hwh+rZBLPxofNpO9tbdBluj8up2uW1MWzMWjTxxbvSxvb/SqAQA1+qT96jnTBx63m4YMujlgnE79ULNGynHRZsJfjgzYL0WfuLqTm1wSwtMTbsFrog7/JeI9y2I+YpFy1wCxa5I0vZMH6eaA7nDefULn0V68PhZVzADaGJYjlUf8LOBJvJsPGbj3++8eZ6DtnB6JvF+JBRQl8fPaieQ8bfNTZAfM9QrFTvuG3JwqzluzA2r0XcS1LrI0HFfAQmf25oxU+6euB4AvpKDXRxqvKvTLsCzyGTv3t8M9WMh+/m47v+iyH+94ruJtXs9+hB6PQa9hS/EPG4r9/mo4vuy/Fr85HcCuj+lpFlmuXrsHeYRM+azMPH7aYifdkTnwjC82+c6kQMawqFUX52LsnHM26WeMDGbc/NZqF7/q5YPuJJGQbI7zIQWhIBNr2XIBPZLHoOWW9ukBmvsMOtO2/RMnqD6Jo7j+fivIaCP8Ql2LiMXL8csG0efivH/7AP1tbY/a6SMQk18SahPib+GOGu/DACj8MWIEpSzWrduJ8X7Qf7IB3mszEDwNXwmHrGWSVPXxG92XDlBcC+LQ7aVjj4Y+Pms7CZ10dYS1MyCjQJ9FjVMiqfPHaHSRn5CNPQHah9Rq1mvaauxvXs++Lwa8V+n9PHI7CoFGO+LjtQrju5WpZLShl+dnw8ta08q6/ewuYOKBVbzsMtgnExVs1V8eoyAsYM9FNBMoWbgEXRJDycTj4KH4RDe0LEfDZHhFIrdJea5aK8lK4rNiOrztY47NerjgQc7dKK3gs4E+/Pb1KdK/s9gtCGwH4H4d4IOAMNRdtNtDnmC1AaL3AC5//MhvDHUIRnWgAKlUe4tCBKPQb7ohmIzzQeYQLugxZisbDViPgbCqM9bWkG0lwWb4df5UJO3ZpEE7fNKrnoSyucfHoLO9+09sF8zwjqtxAvC3rgSyaLOnJaVjvvRdfi7B+128FFq2PQo50Spe1woJCrHHfjo4jXPFjfxcMF95923kRhtnsxunUUrG8tCfvV1YiXjSiPiOc8VlHO8zzjkRqtgbihXn5CNkXgb4jHPAPmTitxm9ExKV09bcnyyOZeMcxYLQT/s+PszBAwPDg2RTD36Rb9yqQdjcLJy/dQYHhDlz6eiMORWLkZA98LZrY8Dkb8U030fCEfxsOxstzBpl79BD3cjPg6LwZzWUxby8LR1OxBlqOWI0pbseQo2uNjx+iNDcbtvab0aqfaHPjPNGhjx3aDF2BSe5hyOMgPyhGWPARNOlqhWaiSXoGXdHefXQfwYERytr5z5/nwm7TaaRkFSHrZgL6j3LAd32XY8bakyiqqAYb+ptjL1zGeFGE3vnFCk7bo5GeU4jbV66iR/+F+KDDEkxaEY58A5A9evgAxcWlKC6txP1aHcVPliuyUE+d6Y732y3GjNWRSCt+iMKUW5gxexU+aDkf3WYFIN5IISoq4j4W9zoeyQKag7nzvfBDZxv0mrMTV+7WBGC9cK8h4dIVWNltxBcdF2HAjA1i1YhW3Ess2Q2nkJxpZMk+LMGO7fvRYaBYP2PWifWwXHjsit7SjsvJugtWnntQjo1ieXQd7Ii2ozzQdYj8HLoM3eb442JKTQ373IkLmGXljc86L8HASR7oMNQZ3/ddBvegq7idU40Z5fl52CCW+A/t5qLH9K3YfzZZ+0NlKdzcd+DHTqKQ9nKBf9RtFBgtINkpqdiwfp8oX7bo+dsq9Ph1BX7ovhhjnEMReUWz1rTyGNGnYjFopMy/Ps5Y6BuFbPk8twIKcnIQFBSB7kPs8d8/z0a3P7bixPVcs26pl11eAODv4/ixs+g9eAneEQ1yjEMILqQUm2xcPFYXZz96cA+pt5IwfPRS/Chm4QSXUKSLKacrCOVlFWIS7RcN2x6fdHPG7pNJyDMy9Y4ciMCQCe74sqcLNgWfh73LNhEGmcBT/RBeg+kPsT/4OHoJCH09YBXct52Af8AxYfQS/NTPFfPXRIjpXP6EP1wv5WVFWOK4BZ+3mY/3OzvCY19cDa1A3xgqLiqCi+MGNO5ui/aTxDy8VYRigzVyr/I+Ys9dEm3OHe/9sgDWvidxI71I/Y0l9cZNWNn64lMxLW19T8BjfTAmz1yJz6Vvngevw+D+VSVWNInZ1t74c4uFmLXqCK6mFKkJxlJWkK9Ar3m3BWg6whPLd1+qMRFUqSyBv99BdOxnJ9qbPWx9ImWCF1c/d78CcecvoaNoYh0nrMcCMVm3bA7EL32XoOPkTdgbkyNWljag5aVlOH4kSha1pfiq5zJ47L2EvGLDQimAWZGXJYvRVnzXwQpf9XNHYFQiDCypLgLAJaLpLly8QbSj+fi2vxt2ywQrNHFTPpQFqlz4qGs8eRmZmD57NdoMdxOrai8CQs6gyzAX/CQTy10W8dwibWI/vH8fqQnXxdLzQOOe9pjpcQi//rFOFlF3DLUOECtLWz5LS0pxcF8Yusri2nb0ajj4HEEnURraDXfH4o2nUC4TsSw7HRu8dskiPQtD7IJw+HK1y/CIKCO9h9rjv5osgNueWAH4Qpw/dgqtui+URXs13EPipY5q9ZfdiAiLFBlejvc6LoXPEelzSRluXryIjr3EImm1UJSVfYhJMbZ6tEun61/uIyr8NPoNscOXfVdg2Z5LKJPXc66LdTbVVbTJWWrhPRST9oTl+ujBfdy9fRsjJskc62SH8Y4HqnhlWh5XlMFj5XZ0Hu6C9hM3wT8kWhYuL/zY1RYzPMJx8662gNDdWZ6diRWu29Gkuw1G2wfi9wWbMXCcO9qO88GJeG3e0oKMOxeDMX+swc99nWG7NgzDJq5E+yHL8OvSYLFgdJ4IlpQWwmnZNrTu74C+ohhs3xOBmTa++LiDHeZvOIsrItt6yUxNhZMoRx+KwvHrkmCcoUtMSm5aOuZar8f33Rah55ztOJdY0xV8MFiw5jdXfDvQA967TsDdKxCNutqg++zdCDG2xMXioXXSoosVmgxZBa+gWMMftJJ8Ow2OjpvxkVjELUQmdp1KVpbeqy7PDfAluVli+gXi0xZiIg1YKRrOJRGxWoow47ICksX4rL0wa7Iv3P2Ow2vHCXXllOeWo/hthie+7bYEzUZ7K3+a7nYrzM7BvAXr1BV3w5fsV77uPQFh6D3SAT+OXIc9p+8YvivSLIDmt3U/2vZZhO8HrcTw2Rsw8g9PAQIHWHtF4PytmiaWablXUQ6/bQKIA+3xP43novMkX6zccQaxt7JRer96suVmZmLabE+l6f88aAVsvMOx2i9S68vWCNg4+6NZ36X4vLsT1sgikZanoTa1022bgtBNgKXZKC+cvZGN8+disdDOG++2WwKHbRdwO7eai1HHz2Ls78vxjgBAzykbsHRdGNb5H4f3zih4bAjFbBHuz1vPQc8Z28RENWgohnK/vEyZqP1HL8NHbRZirGMITsXXvEQ69XYy1qzeCd6UNGv1UUTfyMKFqNNoP2gpmg73xIbwJBSXa0JZVFCEXVsD0aSXHRrL33ZGpaDI8DetiBa2IRDNeizEpz1csOtYAkz2pHCvvByxJ8+h/0hHvC8m++CF+3Ajo3ZXDktZYSGOHTqBpj1s0Vcm5C5Z/NNvJWOsgNH3Ii/WAsgZhs1pbjpHR51Hn5HL8XOfZfAKPA9Hp23oN9oVnadsxk3Dfs3thFuYNM1NLIBVmOwchH2HzqB1P3sB+7Uix/EgNidfT4Tdko14V4BxjEMwDpxPQWZmPuIu34DV4o1qIfthyBqE8Pc5BdjldwBfi2LQedJ6BJ5LqYrEUUU08p2ycHYXa+vrgasQFJMNrmmpN29g4tQV+LDlXFFelmH6ioM4eCYRmSb+6fqUR2J9Buw8gqad5qP5r17YcjRB/b48+67wYDO+aTcP77dfjNG2AfALuyKWbbXScU9k5fKZaHQf5ogvRWYX+JwSa6emv5nlgVhx1y9dFUvFGW1GrYGLKBWZGdlwEYXr5y7WGCFK3qVkzbXKCKQ7wsNZ873xTUcbLPI9jpVrgzDxj9X4Sebm/gt31XMloqgsc96kLkXvM2sbQiJiMVQAvnn/ZZizNhLZhnbck29HR0Wj72gXtBi2CpvCE5GakgK3VTuEf1aYuPwILhhZuNcvX4XVovX4e+PZ6DjeG4vWHILvjuOwXbYTXUa5oe/0jdgWkYCcEk1J4YJUkH4X1naiuPVeisnuEbiZnofDYafQsrs1Wo3bgB3HbqhnWSoLcwUr9uPzVnPQcdJm7D5+0/AXreTnFmK91x581XI2mgxyx7bjt1BiZNW9qvLcAE+/rZWND/7ceBZ6zvRH6LmaAGNcqG0eE22zmWjo77eywnc9xGQT05kXZ3NTonk/R3zVYQE+62qPAVa7cfWuNukrS4pxUrTGVr0Xo42YeFuOJyntI+zQSfT71Rkf91qBTYeuaW4U0XZKsjKxUsyvnzvOQ9NBy8WMX4Qvuy5S2ui5xJo+fXPloQjl9SvXscTJD991tMZfpG/f9XVREy9KwFHTZh/KxEzEkInuakPyUzFTuXnDfnDDkT9/6mGHD9pao4UAxu6Td8Qa0d5Lv5OM0RNXoHF/qXP1cWWysa5FSzfiL82tMU8sjKtiBWnlAYIDj4iFtEi5rb7uKotff8M3ZLFrJlrM96I1vdN8rmhcwcoE1Mv9ynLExVzFkDHO+LD1ArEyNuHYlSyUGS1SbM/B/ZHKVdRomGgYUbeUpnntYiw6iqn8g1g8K3fHotDgt87JysUKRx/82MUGXaZtw/HEEqP6pH/lBfD03CnamrS1v3kNvjC/CP4bg9GsywJ8wc1pWdCynwJmVy8lwHqhFz4VC8R6w2lky3x8IFbAHLF6fuxii7Gu4bidrQE3w/GC94WjlchVk6GrEXIqEds3BmLkBFc0GbUO5+8Wo6ykCAdEQfim7TwMF2vA7/AlHAk9hkY9FqHT71uw52ymkrELZ+Pw+6y1+EuTOeg11RdLvcPgtSUc0618lUb3hcjqLBmvZFm8szKzsXyFn5LtwXP9ES1A88AoOuNRRSlcBWRb9rRDm4kbcS5ZG+PiwgLs8D+EnmJx0l/7z9YL0Xf2NmwJu4qs4nvP5LMtysmB57p9+OgXWfCnb8WBM7cNf7mPI2Fn8JsA60eykPyt+Xy0HbsWLttPI8mw4JUWFuHwnlC07L0IPwxehVX7ryPfTPRORmoG1q3Zie8622Howr24lCHg+6gM69fuQjMZ906zdyLSsLH78MFDXDhzEcMnrcLnHZfAO/iSclvMX7gOn3Vfhu1RySisuIdrMXHo0NcWncWCXLEzGjHRseg1SqyzAW5w3XMF+Qb3W2Z6FhzsN6JRLwcMs9mL1OJHeFhRhLVeAfi85RwMk/acvFLtFjx97LRYxh74e/N5+FpdcemMJqJ0fSA8+ETG7vdl+3GD1jwHW0pFWZlYdeHoOsxJBXCEXMhSchAVdQGd+tgISK+CT0ic9rCU7KTbgjXb8c9f5mKIbRDCYmq6JLMycsTS2SEW4Fz8MnwNAk6nvF0a/KkIEZopK/D3ZnMxwfkQTovWXVvJErNo++YgZcZ918sRfaf5qnsFJ9v6YaL1Foz4Yy2+bTsX3/RwwO+uh5BpEK7kxDuYN98TX3ZerEzzMzfzVHzv3sAIAXgX/LXNEqwUEz2nXIbi8SOkXL+JWVbr8VGLeRg8az26/eqKRr3t0ea39Qqg61tu3UqD75ZQdB/hhH80nYkPxAQctTgI8Wmi9Ty6h/joi+gxbCm+6roYbX/1xCRDX0i/zduAVqLlftpuIbpM3SRAWKwsjKLcPNHyQvCDAGS7cd7YEn5dxddejL6MubYb8adGczHReb9o9YY9hYpieMmEbdTJCt/3ccKgP9Zjmp2/xrdFfhg61VP5mf/eyhZzPY7genq1Jnw59hqmz1mNv4v22WHSBmyPSES5IX6fhaZ/xu0k2IqG+mlrK4wWEzYiNhX5efk4fuyc2uP4pqejciHlFWsAnJacigmT3fGVLGijFgciqeRh1R7K44cPkZ+WggUyeT9vPRctx2/GUTM++Owc0dZW7MAXoun+LGbtVln8CmtYATXLo7Ii7Nh+EC0EPFqP9YHvAZnwhWVITriJOfNFE+y+CF1n70XMHU0bLSkqwSYfAZt+S9F+8lacTSrCscPHMXGGB77s544DsRkIC4/GzJmr8GGnpfAJjce1xDRsWbMN33SyxTC7IFzI0DS6Y0dOYfQksZ5kAeWN9czS96m6ccwOHcd5wW59BBIMPL8h7Zk1z0PkZLHMhVBcz6o0Csl8jLL8XEyb6yULsh0Gzt2B5KrIr8coLynBobBzmDLfB1+3ny+Wo4zZxPVYf/CqWI0mJlAdJf3WLTg4b8E/BNBGLgpGZFz1hj0tp5iL12ArikvzXotUIMOXvZ2xIiAG6UX3UZQni/NKf3zfaSHaCdDuVwEQpt++jzNiHdE19WP/FbD2OYGUvFLkZWaIFbgDv4jV/MNIH+w7m6qe5r5a+P6jCqy/6O2GfdF3cfF8HJyXbcZ7Yq2u2X8VUdHX4enmhw/b2oq2HoFLt7IQGRyGtv3t0XKsNwLkHaVEPCjDycizaNvbFs0FfO03nVYx8qnJaXAWefquw3x0n74Fhy5UK5m7/Q+jn1gaH3WwRb9p6zHJ1l/m5iZ0H70cX3e2FWvNE6tEgckx7Mdlp6ViynR3dR/0kIV7EJtUgIKiUgSHRKLHkCX4VizCZf7nUGRYEC5Hx2GOlZeSjVnS9ovJhdofDCU1OQXzrDzxXnMrdBZZjEx4y3zwRw5GYthYJ5kA8zB7TaQwpGYHjcvNa7fhJFrqJ8KMQfN3wD/yJqLjkhEbn4rzMTewd+9RNO5kje96u8B20ykRbGHEvTIcCI5Ai25WSov6aYArflu4XYFof4Z4ifb2TksbOGxj/Pd9wfeHuHjqvDLv3mu7CA4bIuC6eg+G/uYqE88eLjvO43bOk2anuUIfcPrdLBw5chq/TnYTE3AOvhSt9ODFNBTkF+L4gQg0aj8PrURIlm47h/OqLym4knAHp87EYbCYkV8KEA5asAt3pG0sV2Li8dvvrvibTK6POy3GoJkb8bv0ZeSMdWghgPTXxnMwUoAzPE4DxtKsdCxy2IJP2i5A75l+8A+9hPgbd1UioasCSpu3h6LvCHv8s7Mjlmw5C31/6Ub8DbV5+IVoqM1HrcXawNgnQtceP6jE3l2H0W2wPf73d3+g2VB3jJq7GVNst2LwZE981sYKXwvAT5eFI9vg375x5Rp6ibb/dQ9HzJbfG8d2kPcM5+wxxB7/+MUK41zCxBJ5Uh4yxQpY5LRVNMn5yro5cDW7ysdvrlyKvoTpcz3xlx+n4dOuDug9hYvcdoyb7YNW/XgHrC3aTdyCM9c166wgNxd2C9eqMLVhi/chvfQhYs/HYrq1j4ogWeQZps42tB+wFP2sA2VhKEbS7TTMn7tG6lqMWavDkaWG6wH8t4ag28DF+LKXM0bb+GPBymDYCLltPo7gyATcTCsw7OU8xpnj5zB42BJ83ddVrJJzyBdsrFK+Hz1A9p0kDB0vwNJ9KSYsD0OOgad64YZqdHS8WAE78W2bOfio0xKMdz2MFOPwkaeUWJH9P2atVv78OetO4pL0zbjQxZFwIwVe3sHoxOivxrMxYFEIom8XID8nDza2Pviygw36z/FHXHoZTLHobtIdrPLYgY+azcA/Wtug7RhPTJax+H3BRnQVRegrWRy+6bcS/sfEcpPn79+7h83eAWg3wAHNZHE+n1aKO0kpWL1mFz6Q96c5BcHKcSf6jXRGq982YO+pZOTmFWL9an80Esu7j8h8dKq2EN5JvK0tDE2m44N2ou3/5iWySmVqI9oNclThv23Gr8eeU7rVck/aukssksXKkgs4dgVxiZm4EHsL4YdOYOBYF3zSXubnvECkZIkkV5YibP8xNO22AH9rJlgzULDGZrt8Yxv6jV+J7zsvxOfdHWC38STSDHrUsbDTGDneFf8UfrvticGdquASKQ8qcOLYGbFMbPDPdrLouxxSOGAm5P6ll+cG+KOhkRghAP83Ad9JIrTnRLuuWbQNVpYL0Vfx+9SVeLfFAsxfF4l0Y7kVZsSdjhaz1xZNhq3B2v3xSsNkhMjCRevxXos5KnKi2cBl+GWAo8px3Fi08q862oi5JfV5R+JyRoWKXgkNDFMm1ic9XbAjMgnnz12Gi9Mm/PmnGegxaweCztTmRnqMR6JxGZvVqjwshZfnDvzQcR7+2n4pAkSAku5kwH9TED5oNBP9Zu/AkavGETJAWV4mRoxfgW9EICa5hKJS6r1XlI9tW/bjO9HQPhMN8CfRyFsOElOwvyOa9HPAt7JYfSD97DtvBwLPcCPnMVLj4zFl9hoV5zvN4zhuZehan1YiBFD7DVmEbwZ5YF3oNRAyCnJzsMLdHy162uKLHk5YsSsGd/MNG6F6efwQmXdS1B7Ctx0YT74Yvwx0VkTeNumzFJ+1nqesqdEOIUhRB5fu4fyJM2jV2waNZcK47ryg1SWFm2Q3btzBQtv1or3Pxzd93eAbdh05+gasUckSgF/stAUft5yHRkM9sOOsmK0mSMJxeCzWWGV5qYBBANoKkL/XegGa9HdGC7axr4y/tPNbkZfPZBFtMXItjhu01ay0NIweu0xtvs5ac1RZTrdEu15gv1VAxRqdhq/Aj93s8Mswd/geSQTXvStx1zHo12WyoLlgmd9ZZZY/rizGcpdtYj1Zo+1EAZ+TN5Ei1lZqpmh1JoevKL9h+yLQofN8fC9mvFtQnAI4vVQUF+FE2Am07mOLnwa5YVlAXNUp4RrSJgtB3NlL6NpvEd5taY2+NvtwK19AQf50/959lFXcrzpxaq4cDDyKQaKxfiiLw6rgq0jON/MNKTfiEjBD5OrvP/2BdjN24fiVDGSk3sWEqW74sutSTHQKrVIWqorITFDgMQwRYHz3l/n4obcDWgxapsaiidCPPRare2S/6u6MDQeuooT6WUUZliz2RbPeS9BbFo00WWzzc3KxYUMIvmy3AO2GLkezvg74Ueqy2XQGd0RO87NzMNdqLb7pshjjlwYjQ8UvPkCQ8JcRKXR7Mna/lXyb320ickA35ZdiWf00zBMbj9zQrMqKfFjbrZe/LUav2buRXWUlPkbOzQT8OskVf282H12n78adTFnkExJhtcBL3Wz3RRc7deCumZqbjvi552J821mwpuMSzFx9DNeyNb7u3hUmCtJSfNrFCbuiblcpPPdkYYu7eBWL7Tfh3UYz0OJXb6wTa8VEal5ZeW6AP3/yPCbPWIk//zwTHSdvlk4mVgnTo/v3UFBQhFupeSIb93Hs6Dn0kVX+/U6OWLb9HIqF3/qz2emZ8NsUiM/aL0CXaVsRckE02If3sNE3SB0K+arPcsx1P4h1O0/Ca1sEvLZHwskjECN+dxdNcA7GOR7EiRuFMj8ewNfTH63FjGskC0VYfJEy2Q8Hh+OHtrPxafdlsBPrgCFRT0yTh0yzUIrElAKj6IdHuFdcgOWu28R0XYDP+63C0Ut3kZCQBGcxhf/SeD7GLglCTHI+3f+qPCgrQezp82Ji2uHHgW6w23xO+fjORJzDuN/d8K6s9mPsArBiSyS8tx/HWr/jWLPxMGbabMD37eeg9QRfbBBw5LfPH43C8LHOsiAshsPOy7idbbA+5GP3SwqwznMPfmg9B+2nbMGRuLuolAm1d8dBtGC8cwc7TFx+ELcMh2S4+N0Tk5mlQLS1oJ2haNLNBj/2c8E0eW6N8NXbPxJefsewZsNBDBaQZGhYlxk7cEW+W1FaiEN7QvFju9liwm+Ez6EbwkOGwZbIIhqHpcu24ctWc/Bh+yUCEAeRlFthVlvJzSnAypU78U27+aKlLsVcn5NIzdemhjqmLppstpjed9NzkRB3FQN+XS6a4SIMmLMNKzcfxfodJxTPvPyPY8K89Wryfd3HBfujk/FIQCgp4YZo5/ZoPHglnLaeU/VmitbovHwnPm4xG+80mYGPO9tjrH0w0lRs3EPRtE6hQ/9FaDxyHXwPxqt3CjPuYtrcdWI9WasDZDdMDlzVKPcrcCgoEu06Czj0FBnbfEYARZexx7h+5SbmzFuL95vPRc/p23BcrI2yklKRN+lnTlmV7FSUluL0CZGdXjb4REB6otsR5MoKlC8W46X42zh5OVUWzXtPyq4qD+HjG4zWfcWKECs46Fwy8koqUSxzkNEw+gEvKk6x0XGYON1DWZIjnA7h/LW7uBF7BX0G2+HbviswxysKeRWPqtJm0NWSk5KC2Vbe+KaDNdqP94GT92Fs3BUFT5GbdSI38xy2o+NQJ3zU3laUivPIlnaXFearMwY/9nTAb0tDlMvnfnEhdu8IQ+MuC/Be05nqnEKH3zfgzC0tZPJ24i2MnLAM38hiu8A7Si3QGbcSMX/henzcygq9pm+Fvddh+O48ob69fmcErBy2osNAe3zY3QXLZfEsrHyMsrtJ6szMV93sMXxpKO4YIoIyUzOwRXClebeF+KybI6w2nkNaejYCxJpt2sVaKR0zXYKxcbehbyJvTh77MFKw5p1WNhixJARnk+kOrIC3WCctetnh+6FiiZ5PU/tN5aXFiD4fhzmLfPGNKHPvtLDFgnUncDm1ekP7VZfnBvi89DR4rg1Qh17+LhpHzz82YfnWSGzaEwU7F3/MtPeH+55YZGTmyaAeEpCdiU/6uMNbVjNjIb0RL6vn3FX4n5+nY4jdPpxKLFCuln4jHESTWYAJrkdwPiEThUVlyMzR8nwkxt9U8fd/+3GKaFhbEHAyGffLizFj5ip81HwOmo9Zj9hMTYO8djkBYya44P98Nw0tx/lg77nUGv5olsdlxQg7fBYjZ/lijlsIVvtHiRBFwtrJDz8LEH4mE27yijCk5ZXh/NlYjBrrgP+vsRX+cD+K2znV+WF4gGf3liB81GQKmohm6XUoUbSjdFhZr8Onshg1/80XRy8koaC4HFm5zC9TjOK8XISLNdSs82y838VZuVt47H2b7z406SgToe0SbItKRo4hzl7zn9/G7AXe+M9vp2HgwkCcj09DgvSz28Al+NP30/CX5vPRZrQHJtj4YeoiP4xfsAUT7HaKqRqPgANnMeJXR7zbxlYWx/24mqbly8mWtuQK6GRm58HRmVEiAsI93XHyRi6SbqfCc+V2/P2HKfhATNuuk3wwWczXsXN80GGIkzpZ+nVP0ZrdwhB1OdPs8W8WhlqeCj+JPsMd8D8/TRdt3xVzVh6AT8BJrN16BJOtt8DR5wg2yARzdt6Mj1tbobWMWaCMWUpmobSvGNl5JSIDJWIRhaJz38X4r0ZzsSY0ASnZ+Th99AQ++2WmgLUXPA9woZSpmJ0Jr9UC8I0m4//6YjLaygK1NTxRuSAeFOZgw7rd+KjRFLSf5od959LkjYdIuHgBvUX+/tyU0RnhSDI5lFOzPMKF0zEY9auDWLOz0XjIKizwOIwN0icP3wMY84cnPmkxF02HrsGavTEqx8stAVQXtwCMmL0Zy7dEYOOeU3DxDMbA8W74R7O56D51izrBybMXQXuOYvQ0TxnnABxPyDGjCQoYl+TBxs5HZGwevh3oiQtJ+UhPuYudYjWOnLsF1qtD4bvnNNZuCcfIqatF47XGV72WYcvRm0hOyZIF6hi+aSnz+BcrNB2+Cr9Z+2Hc3I2Y5bJPpchwctuBn7ssxDe9l2F1yBVclwWvoFBkhTJTWIajYWcxcbI7/rfI4xSPY4hNzkNG4nV07GuLT7s6YYbnSS0f1YNShAUfQ/NOc/F/fzYRn/dyxeLNZ5HHKInHFThx5ASadJiF74eshvu+OKW0eK7yx8+drfGpgL5fRCIS0wtVjidiQWFJOY6En8WvMr//n29mYbpYutczS3DlrIzf8KUKU77r74pf52+RebANQ6d4qjj/L7o6YLxjCKJv5SIs9ASGimXyp8ZzZXGLxOn4DMGaUlV/XmE5rsoC7bFqO/7zq6lo/dsGBEWnyZzNga1YJx82nY5/tLNDn6m++H3Rdoydu14d+PukrRV+6OuCOe5HEHMj+7W4ZvTy3ADPgzbnoy9jpvV6NJaV7Mtui9F86Ar0mrgGbYcsw8Dp1PQSkJGRi8CAI+giJtbghXtw6CInUXVJiL+FhQu9xfx2gdO2M0jMKsG+nYfQe7Qrukz0FYbexRPRReVFChQZkz1o/m4En7yFyqJ8zLH2RZuhbpi47CAyDH7nPDEL/baEoHW/Jeg4wQfrBAyYpKxGqShB6MGT6DVqmYpQaTfKHV3HrkSzgQ5oOsgVE+z34Zho7wTX8zKZJ0xzxy9jvLAm6LJoO4Y6pGTdzcSOzSHK9zfBIRhHr2QjPiYek2Z44pfBK5S//gkTX8qdGzcwdupKtBq9Fm47olWytM2b9qP7KDd0+X2jaDh5VWGjbEPy9ZuwcdyGn/otw1JZEG4lZyH2bBw6DXZGoz4O6jh61zGr0GmkG9rIWLQc4qLSCKzeIVbQ1qMCsE7oLdrQ9mM1Q7v0snfvEQyY4I6WYl6evZ6NmzdS4LNuH5qLSc1ohHZSb1sxsdsOExq+An2nrofDhhM4ey3rybEyKmozNiMD3uuD0FN4/XlHG/zUfxk6SFt7TFiNtiNWwUbAyHt7BGbM80L70auxyFc0SjP27YWTMZg91xs/93GGjwB2Qko2IsOOo/0gR4wT3h+KM2z6i6zsDwwXC9IBTQaugP3GU7glizJL4d27WO8ViF/EEpjteQznuVEmFmfc6bMY+8dqtBq5BisDrzw1v07m3XRsERnrPmIZvu2yGI2lT93GrULH4S5oOcAJHUd7Yvm2s7gui6lwAbcv0f24CY3ku8wJ00WebTt8ORr1d0LPKTJvgi4hs7ASD4vy1Nz4QLTdpuM24OhVLbKjRuEZBB7uWrYFHYa7YoTdfnWKNk/atNl7r3zfUWTaBZ3HijyMdkMjGUPKgo1XBG5nlSn/+/59x9FpkJOKbGst49l2mKuKbOsz2Qe/2+/GZCtfdBy1EhOWBqlTr6YlLTEJq8Uy+7mnPax8TyH6ZhaSrl7BoHFu6P3HFvgcvmEIG72PS+diMG6Ku4ybE8aLZn8qIVeLFiorwtEDkeg6aKlKMRJ87g4KsrJFOfJBe5Hn35wOIMX4RJKhMJ2Ei6sfvu9qj8UiK/FphTh/6iLGz1yr0p60lr6wP22HuigedxyzWhSRQwi/mIqKe/fgJ/O13xjp82/rEXYpA4zXMC7lxUU4fPC4kpGhc/1FWxdLKiNLFuidaNPfHr8MWY42Uj/TgbD+1jIvBgj2OW6IxMUbObhf26GbV1SeH+Cl3KuoUJt6bl77MFS0jHYjVqDzuDUYY70T3oEXcSe/HJXlZYg+dwWr1odi57HrVSFtermbloWAPSfEBD+OM5fTUCKr8qHQs1i14Rj8D11FYYVRPhRDYS6XhGu34bExDBuCYxAnAlUpA+EfeBKr/E4g5PStquxyjN29e/sONviFw10sjNDoO1UZIavK40diHqZg05aDGDfLGx1HyoD96o4h8zfD3f80YhL107KPcePaLWzeHg7PPdE4LxqAMZ4V5BXi5IlYLF93CCEnriNDrI34SzewbvNReG4/jYS0ElQ+MeCPVZid/94orNwWhaPRSSqeOzLqkrxzHOtDYpBdXG0lEOAz09IRdPAMXDcew6krmSgoKMEdab/npiNw2yC0MRyu6w/DWdrhsPYAnNaFYuXWEwg/eQ2HwmPh6hWKfZE3cMdkLLTyCJcu3xQzNVIdFEvKEEsjIw9RUXFY4RuGZd6H4CLkuDYUbps5RjFiYaWj7IlTTbWXTFkIg/ZHYarNJnQVPrcd6Y7+MzZh0boIHI++hZjYRPhsPgLfwAsyScxHZ6Ul30XIgdNYvj4MJ6/lKg3/Smw8Vm4IQ7DwPq3AsAfw+AGuJ9zCRv+j8Np1FjE3q+sryMrDMeHH8rWHERadjOyS+4KXD5B87Qa2ivm/ZsdpBfp1bQSzcJM5NzsXu/dGYobNFnSTPnHC957khbnLgxEUcQO5hg1TuqJK8gpw6NA5zF+yDb3HuqPNCFcBdm/MXnUQh84xR5NI1aMHuBEnIDl2OT5ssxBTxTq6aRS7XlXk2zy1GXLgFFZvjcDuE1q89cOKcly6EI+lrrswaNJqac9ydP5tNSY57MX2w1eRZjgpXl5cgosXrsND+EaZUXIj8rFk9QF4bjsO/6CzWO8XgXU7zyD8/J0n5iJLYW4eTp0UufI5rJLEJaYXqE3ZdaJMbN1/UazEwqqooqz0TATsO4FVIo9HZC7qYlNZUiLtvY6V3ocRKLKZeLcQOXezsEEsu7XbT+L4pXSjyKTqkpOTj6ORMXARGQ89dRN3xcq+nXgH2wVT+DtXkQ+ntQdlLoTC0/8EDvIZQxqJx/cqcGD/KcGRo/ALi0eBGe2kvLwc8YI1xK/NQefVCVyG+4aFX8RKH6lf5gT55STkJvPc/3CsyFgWigwK5usuLwTwgjbKnUCTJi0zH7dScpCUJuZZTjGKRFPlphABqbLinnKxlJTff2Ijkyk9S0orUFhcgcp7D9QmW1kZ/18uz9cSCyx1MiUsn+Fxbm7mqonDekoq1IZUVZFnuRFYLAsHn1cpPM3USV8jU4tmZOUjKTUHt1JzkSZmWkGJ1G90aIUbXqyrSL5bKW03LjzcQe27QPrKNlAgebK1SL5bJO2qNWWtaLY6D5gelTwrF57xneIy45A7rbA/pcIjunq4WDGXD9vP77CPJLaBxLDCfPmp6pY6y8qr2/fEprKh3GMfDe2hL5btq5B3+b2qevlT/s8x0jfT61tUCllpf7YsiMlpwmuRGwI0U0FUCL8q5fvsiz625grTIrOOAr4jPGA7mdJAjbGB93qhrOhjZlwf28H0x+wP+a7eEd4/0MdY+K9S4JpnU43CMSgVnjHVbTLlh33KKFDpcdke4yHks/wu00DfUf3Pxp30fOXu0N2HPODn4rxZxW53m74Fp8SSKqslzI4LDOWBfS/VvyWNJj+Y8jYtPU+157bMzcy8EjUPdf5w3lQKz41lRo2vtIW/o1xq8ltZI8WxcdHkXuRD3uP3mV5Bl0fKr3G6BcoS+cR5alwf28G5QpzQ62AsvT53OMbmCs+ulBvNOco0U2yz3VV9MRDrKZNvVKUpFx6otrCfwhPjMdIL5+I9hXFlmjxKuzh+nEf8HeVP/0bVfJB2m6nqtZQXA3hLsRRLafByv7QEJyOi0WuEC4bM88PGI9eeiDZ6G8ubAnr/PgX4/wGWZFU448kxJgAAAABJRU5ErkJggg=="

st.set_page_config(page_title="Faculty Workload Allocation System", layout="wide", page_icon="\U0001F4DA")

# --- Navy blue & white theme -------------------------------------------------
NAVY = "#0F2350"
NAVY_LIGHT = "#1B3A6B"
st.markdown(
    f"""
    <style>
        .stApp {{ background-color: #FFFFFF; }}
        section[data-testid="stSidebar"] {{
            background-color: {NAVY};
        }}
        section[data-testid="stSidebar"] * {{
            color: #FFFFFF !important;
        }}
        section[data-testid="stSidebar"] input, section[data-testid="stSidebar"] textarea {{
            color: #000000 !important;
        }}
        h1, h2, h3 {{ color: {NAVY}; }}
        .stButton>button, .stDownloadButton>button {{
            background-color: {NAVY};
            color: #FFFFFF;
            border: none;
            border-radius: 6px;
        }}
        .stButton>button:hover, .stDownloadButton>button:hover {{
            background-color: {NAVY_LIGHT};
            color: #FFFFFF;
        }}
        .stTabs [data-baseweb="tab-list"] {{
            gap: 4px;
        }}
        .stTabs [aria-selected="true"] {{
            background-color: {NAVY};
            color: #FFFFFF !important;
            border-radius: 6px 6px 0 0;
        }}
        div[data-testid="stMetricValue"] {{ color: {NAVY}; }}
        .iu-footer {{
            text-align: center;
            color: {NAVY};
            padding: 18px 0 6px 0;
            border-top: 2px solid {NAVY};
            margin-top: 30px;
            font-size: 0.9rem;
        }}
    </style>
    """,
    unsafe_allow_html=True,
)

# --- Header: title/caption on the left, university logo on the right --------
header_left, header_right = st.columns([4, 1])
with header_left:
    st.title("Faculty Workload Allocation and Course Recommendation System")
    st.caption(
        "Upload last semester's workload and this semester's Scheme(s) of Studies to get an "
        "optimized, explainable course allocation."
    )
with header_right:
    st.markdown(
        f'<img src="data:image/png;base64,{IU_LOGO_B64}" style="float:right; max-width:100%; margin-top:8px;">',
        unsafe_allow_html=True,
    )

with st.sidebar:
    st.header("1. Previous Semester Workload")
    workload_file = st.file_uploader(
        "Faculty workload file (.xlsx or .docx)", type=["xlsx", "xls", "docx"], key="workload"
    )

    st.header("2. Scheme of Studies (upcoming semester)")
    st.caption(
        "Optional — Excel, Word, or PDF accepted. If you skip this, the previous "
        "semester's own course list is used as the candidate pool instead."
    )
    bsai_file = st.file_uploader("BSAI scheme", type=["xlsx", "xls", "docx", "pdf"], key="bsai")
    bscs_file = st.file_uploader("BSCS scheme", type=["xlsx", "xls", "docx", "pdf"], key="bscs")
    bsse_file = st.file_uploader("BSSE scheme", type=["xlsx", "xls", "docx", "pdf"], key="bsse")

    st.header("3. Report label")
    semester_label = st.text_input("Semester label for the PDF report", value="Upcoming Semester")

    run_button = st.button("Run allocation", type="primary", use_container_width=True)

if "result" not in st.session_state:
    st.session_state.result = None

if run_button:
    if workload_file is None:
        st.error("Please upload the previous semester workload file.")
    else:
        scheme_files = {
            "BSAI": (bsai_file, bsai_file.name if bsai_file else ""),
            "BSCS": (bscs_file, bscs_file.name if bscs_file else ""),
            "BSSE": (bsse_file, bsse_file.name if bsse_file else ""),
        }
        try:
            with st.spinner("Processing data and solving the allocation..."):
                result = run_pipeline(workload_file, scheme_files, workload_filename=workload_file.name)
            st.session_state.result = result
            st.session_state.semester_label = semester_label
            st.success(f"Allocation complete — solver status: {result.solver_status}")
        except Exception as e:
            st.session_state.result = None
            st.error(f"Something went wrong while processing: {e}")
            st.code(traceback.format_exc())

result = st.session_state.result

if result is None:
    st.info("Upload your files in the sidebar and click **Run allocation** to get started.")
    with st.expander("Expected file formats"):
        st.markdown(
            """
**Previous semester workload (.xlsx or .docx)** — one row per faculty-course assignment:
`Faculty Name, Course Code, Course Name, Theory/Lab, Semester, Cr.hr, Cn.hr`
(column names are matched flexibly, so close variants work too; a Word table
with a title row above the header, like a workload doc pasted from a department
template, is handled automatically)

**Scheme of Studies (.xlsx, .docx, or .pdf), one file per program, optional** —
one row per course offered: `Course Code, Course Name, Credit Hours, Theory/Lab, Semester`.
If you don't have this yet, skip it — the previous semester's own course list
is used as the candidate pool instead.
            """
        )
else:
    if result.warnings:
        with st.expander(f"⚠️ {len(result.warnings)} data warning(s)", expanded=False):
            for w in result.warnings:
                st.write(f"- {w}")

    tab_summary, tab_history, tab_courses, tab_download = st.tabs(
        ["Faculty Summary", "Teaching History & Recommendations", "Unallocated Courses", "Download Report"]
    )

    with tab_summary:
        st.subheader("Faculty-wise Workload Summary")
        st.dataframe(result.faculty_summary, use_container_width=True, hide_index=True)

        st.subheader("Full Allocation")
        display_cols = ["faculty_name", "course_code", "course_name", "credit_hours",
                         "is_lab", "core_course", "already_taught_previously", "score", "reason"]
        if not result.assignments.empty:
            st.dataframe(result.assignments[display_cols], use_container_width=True, hide_index=True)
        else:
            st.warning("No courses were allocated. Check your constraints and input data.")

    with tab_history:
        st.subheader("Faculty Member → Previous Courses → Recommended Courses")
        st.dataframe(faculty_teaching_map(result), use_container_width=True, hide_index=True)

    with tab_courses:
        st.subheader("Courses That Could Not Be Allocated")
        if result.unallocated_courses.empty:
            st.success("All courses in the uploaded scheme(s) were successfully allocated.")
        else:
            st.dataframe(result.unallocated_courses, use_container_width=True, hide_index=True)

    with tab_download:
        st.subheader("Download the PDF Report")
        pdf_bytes = generate_pdf_report(result, semester_label=st.session_state.get("semester_label", "Upcoming Semester"))
        st.download_button("Download Faculty Workload Report (PDF)", data=pdf_bytes,
                            file_name="faculty_workload_report.pdf", mime="application/pdf",
                            use_container_width=True)

        st.subheader("Download Allocation as Excel")
        excel_buf = io.BytesIO()
        with pd.ExcelWriter(excel_buf, engine="openpyxl") as writer:
            result.assignments.to_excel(writer, sheet_name="Assignments", index=False)
            result.faculty_summary.to_excel(writer, sheet_name="Faculty Summary", index=False)
            result.unallocated_courses.to_excel(writer, sheet_name="Unallocated", index=False)
        st.download_button("Download Allocation (Excel)", data=excel_buf.getvalue(),
                            file_name="faculty_workload_allocation.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            use_container_width=True)

# --- Footer -------------------------------------------------------------
st.markdown(
    """
    <div class="iu-footer">
        Created by <b>Abdul Baqi Malik</b> (Lecturer, BSAI Program)<br>
        Department of Computing and Technology — Iqra University, Chak Shahzad Campus, Islamabad
    </div>
    """,
    unsafe_allow_html=True,
)
